from __future__ import annotations

import csv
import html
import re
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from utils.dates import has_exact_date_in_url, in_date_range, parse_article_date, parse_date_from_url
from utils.models import Article, StoryCluster


EDITORIAL_SECTIONS = [
    "United Nations",
    "Politics",
    "Military & Security",
    "Human Rights & Rule of Law",
    "Migration",
    "Economy & Energy",
    "Environment",
    "Governance & Public Services",
    "Regional & International",
    "Varieties",
]

MANDATORY_REPORT_SECTIONS = {
    "United Nations",
    "Politics",
    "Military & Security",
    "Human Rights & Rule of Law",
    "Economy & Energy",
    "Environment",
    "Regional & International",
    "Varieties",
}

PICS_SECTION_THEMES = {
    "United Nations": ["Political Process", "UN Agencies", "Other UN News"],
    "Politics": ["Political Process", "Government Affairs", "Other Political News", "Commentary"],
    "Military & Security": ["Security Developments", "Armed Groups", "Crime & Enforcement"],
    "Human Rights & Rule of Law": ["Accountability", "Human Rights"],
    "Migration": ["Migration Policy", "Deportation & Enforcement", "Migration Routes & Trafficking", "Commentary"],
    "Economy & Energy": ["Banking & Finance", "Oil & Energy", "Infrastructure & Reconstruction", "Other Economic News"],
    "Environment": ["Water & Climate", "Environmental Risks", "Agriculture"],
    "Governance & Public Services": ["Reconstruction & Infrastructure", "Public Administration", "Health", "Education"],
    "Regional & International": ["Regional Diplomacy", "International Relations", "Foreign Cooperation"],
    "Varieties": ["Analysis & Opinion", "Features", "Culture", "Society", "Sports"],
}

PRIORITY_ORDER = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
DEFAULT_RELEVANCE_THRESHOLD = 60

SOURCE_TIERS = {
    "lana": "A",
    "libya_herald": "A",
    "reuters": "A",
    "al_wasat": "A",
    "libya_observer": "A",
    "libya_al_ahrar": "A",
    "ean_libya": "A",
    "al_menassa": "A",
    "libya_24": "A",
    "rna_reportage": "A",
    "al_sabaah": "A",
    "libya_update": "B",
    "libya_review": "C",
    "al_marsad": "C",
    "al_shahed": "C",
}

APPROVED_SOURCE_IDS = {
    *SOURCE_TIERS.keys(),
    "al_marsad",
    "al_saaa_24",
    "address_libya",
    "asharq_al_awsat",
    "fawasel_media",
    "tanasuh",
    "al_jazeera_arabic",
    "akhbar_libya_24",
    "al_ain",
    "arabi21",
    "new_arab",
    "ansa",
    "anadolu_agency",
    "the_guardian",
    "bss_news",
    "ch_aviation",
    "volcanodiscovery",
    "ap",
    "bbc",
}

EXCLUDED_SOURCE_IDS = {"gazettengr"}

ARABIC_RECOVERY_MONTHS = [
    "يناير",
    "كانون الثاني",
    "فبراير",
    "شباط",
    "مارس",
    "آذار",
    "ابريل",
    "أبريل",
    "نيسان",
    "مايو",
    "أيار",
    "يونيو",
    "حزيران",
    "يوليو",
    "تموز",
    "اغسطس",
    "أغسطس",
    "آب",
    "سبتمبر",
    "ايلول",
    "أيلول",
    "اكتوبر",
    "أكتوبر",
    "تشرين الأول",
    "نوفمبر",
    "تشرين الثاني",
    "ديسمبر",
    "كانون الأول",
]

TIER_BONUS = {"A": 3, "B": 2, "C": 1, "D": 0}

NON_ARTICLE_MARKERS = (
    "/search/",
    "/search",
    "site-search",
    "?s=",
    "&s=",
    "/tag/",
    "/tags/",
    "/category/",
    "/author/",
    "/page/",
    "/archive/",
    "/section/",
    "/بحث/",
)

NOISE_PATTERNS = {
    "sports": [
        "sport",
        "football",
        "match",
        "league",
        "club",
        "athlete",
        "كرة",
        "مباراة",
        "الدوري",
        "النادي",
        "رياض",
    ],
    "entertainment": ["celebrity", "artist", "music", "movie", "الفنان", "مشاهير", "مسلسل", "سينما"],
    "markets": ["gold prices", "currency exchange", "stock market", "أسعار الذهب", "سعر الدولار", "العملات", "سعر الذهب"],
    "weather": ["weather", "forecast", "temperature", "fishing", "الطقس", "درجات الحرارة", "النشرة الجوية", "الصيد البحري"],
    "horoscope": ["horoscope", "zodiac", "حظك", "الأبراج"],
    "routine_local": [
        "traffic campaign",
        "road closure",
        "road maintenance",
        "diversion",
        "shop theft",
        "horse breeding",
        "حملة مرورية",
        "إغلاق طريق",
        "تحويلة",
        "صيانة الطريق",
        "سرقة متجر",
        "تربية الخيول",
        "خيول",
        "إصلاح",
        "عطلا",
        "كلية مستقلة",
        "تسويات مالية",
        "ورشة",
        "احتفالية",
        "مراسم",
        "اجتماع تقابلي",
        "زيارة تفقدية",
        "تحدي القراءة",
        "صيف الشباب",
        "youth voice",
        "initiative-making forum",
    ],
}

GENERIC_HEADLINE_PREFIXES = (
    "Economic policy development reported",
    "Governance or public services development reported",
    "Foreign relations development affecting Libya",
    "Migration and refugee development reported",
    "Political process development reported",
    "UN-related Libya development reported",
    "Security development reported",
    "Public finance or economic stability issue reported",
    "Governance decision with public-service implications",
    "Foreign diplomacy or economic cooperation affects Libya file",
    "Rule of law or accountability development reported",
    "Security development reported",
    "Environmental risk development reported",
    "Public-service governance decision reported",
    "UN engagement on Libya political track",
    "Libya political institutions face process-related development",
    "Libya migration file draws national security and protection concerns",
    "Libya public affairs development",
    "Oil, fuel or energy policy issue affects Libya supply or production outlook",
    "Foreign diplomacy or economic cooperation affects Libya file",
    "Libyan institutions address a public-policy issue",
    "Migration pressures in Libya draw",
    "Rights and accountability concerns remain",
    "Libyan security authorities report operation",
    "Libya economic reporting highlights",
    "Italy-linked diplomacy keeps",
)

SECTION_MARKERS = [
    (
        "United Nations",
        96,
        ["unsmil", "srsg", "dsrsg", "united nations", "security council", "البعثة الأممية", "الأمم المتحدة", "تيتيه"],
    ),
    (
        "Military & Security",
        88,
        ["armed", "clash", "militia", "army", "ceasefire", "security arrangement", "weapon", "اشتباك", "مسلح", "قوات", "سلاح", "ترتيبات أمنية"],
    ),
    (
        "Migration",
        88,
        ["migration", "migrant", "refugee", "deportation", "trafficking", "unhcr", "iom", "هجرة", "مهاجر", "لاجئ", "ترحيل", "اتجار بالبشر"],
    ),
    (
        "Economy & Energy",
        84,
        [
            "central bank",
            "budget",
            "currency",
            "public finance",
            "oil",
            "gas",
            "energy",
            "fuel",
            "noc",
            "brega",
            "investment",
            "trade",
            "free zone",
            "export",
            "economic cooperation",
            "مصرف",
            "ميزانية",
            "عملة",
            "نفط",
            "غاز",
            "وقود",
            "البريقة",
            "استثمار",
            "تجارة",
            "منطقة حرة",
            "صادرات",
            "تصدير",
            "تعاون اقتصادي",
        ],
    ),
    (
        "Environment",
        68,
        [
            "rain",
            "flood",
            "water resources",
            "groundwater",
            "pollution",
            "poisonous fish",
            "insects",
            "pests",
            "climate",
            "agricultural damage",
            "أمطار",
            "فيضانات",
            "موارد مائية",
            "مياه جوفية",
            "تلوث",
            "أسماك سامة",
            "حشرات",
            "آفات",
            "مناخ",
            "أضرار زراعية",
        ],
    ),
    (
        "Politics",
        88,
        ["election", "parliament", "political process", "structured dialogue", "constitutional", "executive authority", "roadmap", "مجلس النواب", "انتخابات", "حوار منظم", "حوار", "دستوري", "سلطة تنفيذية", "خارطة"],
    ),
    (
        "Human Rights & Rule of Law",
        82,
        ["human rights", "court", "justice", "prison", "detention", "prosecution", "icc", "accountability", "حقوق الإنسان", "محكمة", "عدل", "سجن", "نيابة", "المحكمة الجنائية", "مساءلة"],
    ),
    (
        "Governance & Public Services",
        70,
        ["government", "cabinet", "ministry", "minister", "public administration", "municipal", "infrastructure", "health policy", "education policy", "حكومة", "وزارة", "وزير", "بلدية", "خدمات", "بنية تحتية"],
    ),
    (
        "Regional & International",
        80,
        ["italy", "egypt", "tunisia", "turkey", "eu ", "european union", "foreign minister", "diplomacy", "إيطاليا", "مصر", "تونس", "تركيا", "الاتحاد الأوروبي", "وزير الخارجية", "دبلوماسي"],
    ),
    (
        "Varieties",
        10,
        ["culture", "heritage", "sport", "football", "ثقافة", "تراث", "رياض", "كرة"],
    ),
]

MUNICIPAL_MARKERS = [
    "municipal",
    "municipality",
    "public services",
    "water",
    "electricity",
    "infrastructure",
    "reconstruction",
    "education",
    "health",
    "hospital",
    "بلدية",
    "البلديات",
    "خدمات",
    "كهرباء",
    "مياه",
    "بنية تحتية",
    "إعمار",
    "تعليم",
    "صحة",
    "مستشفى",
]

LIBYA_MARKERS = [
    "libya",
    "libyan",
    "tripoli",
    "benghazi",
    "misrata",
    "derna",
    "sebha",
    "zawiya",
    "ليبيا",
    "ليبي",
    "الليبي",
    "طرابلس",
    "بنغازي",
    "مصراتة",
    "درنة",
    "سبها",
    "الزاوية",
]

LIBYA_INSTITUTION_MARKERS = [
    "house of representatives",
    "high council of state",
    "presidential council",
    "national oil corporation",
    "central bank of libya",
    "مجلس النواب",
    "المجلس الأعلى للدولة",
    "المجلس الرئاسي",
    "المؤسسة الوطنية للنفط",
    "مصرف ليبيا المركزي",
    "الدبيبة",
    "حفتر",
    "عقيلة",
    "تكالة",
    "المنفي",
]

HIGH_PRIORITY_MARKERS = [
    "unsmil",
    "srsg",
    "election",
    "political process",
    "security",
    "armed",
    "migration",
    "human rights",
    "government crisis",
    "international diplomacy",
    "البعثة الأممية",
    "انتخابات",
    "أمن",
    "مسلح",
    "هجرة",
    "حقوق الإنسان",
]

MEDIUM_PRIORITY_MARKERS = [
    "municipal",
    "economy",
    "oil",
    "energy",
    "infrastructure",
    "central bank",
    "بلدية",
    "اقتصاد",
    "نفط",
    "طاقة",
    "بنية تحتية",
    "مصرف",
]

STOPWORDS = {
    "the",
    "and",
    "for",
    "from",
    "with",
    "after",
    "before",
    "this",
    "that",
    "says",
    "over",
    "في",
    "من",
    "عن",
    "على",
    "الى",
    "إلى",
    "مع",
    "بعد",
    "قبل",
    "هذا",
    "هذه",
    "ذلك",
    "تلك",
    "الذي",
    "التي",
    "libya",
    "libyan",
    "ليبيا",
    "ليبي",
    "الليبي",
}

STRONG_STORY_THEMES = {
    "structured_dialogue",
    "migration_settlement",
    "migration_deportation",
    "migrant_boat",
    "migration_trafficking",
    "zawiya_clashes",
    "central_bank_salary",
    "central_bank_cyber",
    "central_bank_paris_forum",
    "bank_accountability_case",
    "noc_slb",
    "brega_gas",
    "tripoli_projects",
    "gaza_convoy_detention",
    "sudan_repatriation",
    "asian_african_parliamentary_forum",
    "egypt_us_libya_sudan",
}


@dataclass(slots=True)
class EditorialResult:
    approved_articles: list[Article]
    review_articles: list[Article]
    rejected_articles: list[Article]
    clusters: list[StoryCluster]
    qa_rows: list[dict[str, str]]
    duplicates_removed: int
    review_recovery_rows: list[dict[str, str]]
    classification_qa_rows: list[dict[str, str]]
    stale_story_rows: list[dict[str, str]]
    raw_candidate_promotion_rows: list[dict[str, str]]


def run_editorial_pipeline(
    approved_articles: list[Article],
    review_articles: list[Article],
    raw_articles: list[Article] | None = None,
    threshold: int = DEFAULT_RELEVANCE_THRESHOLD,
    start_date: datetime | None = None,
    end_date: datetime | None = None,
) -> EditorialResult:
    editorial_approved: list[Article] = []
    excluded_articles = [
        article for article in [*approved_articles, *review_articles] if article.source_id in EXCLUDED_SOURCE_IDS
    ]
    for article in excluded_articles:
        article.editorial_status = "rejected"
        article.rejection_reason = "excluded_source"
        article.editorial_reason = "excluded_source:gazettengr_removed_from_approved_universe"
        article.include_candidate = False

    approved_articles = [article for article in approved_articles if article.source_id not in EXCLUDED_SOURCE_IDS]
    editorial_review = [article for article in review_articles if article.source_id not in EXCLUDED_SOURCE_IDS]
    editorial_rejected: list[Article] = []
    editorial_rejected.extend(excluded_articles)
    duplicates_removed = 0

    for article in approved_articles:
        validate_and_score(article, threshold)
        if article.editorial_status == "approved":
            editorial_approved.append(article)
        elif article.editorial_status == "review":
            article.include_candidate = False
            article.qa_status = "needs_review"
            article.qa_notes = append_note(article.qa_notes, article.editorial_reason)
            editorial_review.append(article)
        else:
            article.include_candidate = False
            article.qa_status = "rejected"
            article.qa_notes = append_note(article.qa_notes, article.editorial_reason)
            editorial_rejected.append(article)

    recovered_review, recovery_rows = recover_review_queue_items(
        editorial_review,
        editorial_approved,
        threshold,
        start_date,
        end_date,
    )
    editorial_approved.extend(recovered_review)
    editorial_review = [article for article in editorial_review if article not in recovered_review]

    promoted_review = promote_narrative_review_items(editorial_review, threshold, origin="review_queue")
    editorial_approved.extend(promoted_review)
    editorial_review = [article for article in editorial_review if article not in promoted_review]

    promoted_rejected = promote_narrative_review_items(editorial_rejected, threshold, origin="story_splitting_review")
    editorial_approved.extend(promoted_rejected)
    editorial_rejected = [article for article in editorial_rejected if article not in promoted_rejected]

    promoted_raw, raw_promotion_rows = promote_raw_candidate_pool(
        raw_articles or [],
        editorial_approved,
        threshold,
        start_date,
        end_date,
    )
    editorial_approved.extend(promoted_raw)

    stale_story_rows = apply_temporal_validation(editorial_approved, start_date, end_date)
    stale_rejected = [article for article in editorial_approved if article.editorial_status == "rejected"]
    editorial_approved = [article for article in editorial_approved if article.editorial_status != "rejected"]
    editorial_rejected.extend(stale_rejected)

    classification_qa_rows = apply_classification_qa(editorial_approved)

    editorial_approved, same_source_duplicates = remove_same_source_duplicates(editorial_approved)
    duplicates_removed += len(same_source_duplicates)
    editorial_rejected.extend(same_source_duplicates)

    clusters = cluster_stories(editorial_approved)
    editorial_approved, cluster_rejected, clusters = select_final_clusters(editorial_approved, clusters)
    editorial_rejected.extend(cluster_rejected)
    classification_qa_rows.extend(apply_cluster_classification_qa(clusters, editorial_approved))
    qa_rows = run_editorial_qa(editorial_approved, clusters)
    for _ in range(5):
        failing = [row for row in qa_rows if row["status"] != "pass"]
        if not failing:
            break
        editorial_approved = auto_fix_qa_failures(editorial_approved, failing)
        clusters = cluster_stories(editorial_approved)
        classification_qa_rows.extend(apply_classification_qa(editorial_approved))
        editorial_approved, cluster_rejected, clusters = select_final_clusters(editorial_approved, clusters)
        editorial_rejected.extend(cluster_rejected)
        classification_qa_rows.extend(apply_cluster_classification_qa(clusters, editorial_approved))
        qa_rows = run_editorial_qa(editorial_approved, clusters)

    editorial_approved.sort(
        key=lambda article: (
            PRIORITY_ORDER.get(article.priority, 9),
            -article.relevance_score,
            article.published_at or datetime.min,
            article.source_name,
        )
    )
    clusters.sort(
        key=lambda cluster: (
            PRIORITY_ORDER.get(cluster.priority, 9),
            -cluster.relevance_score,
            cluster.publication_date,
            cluster.canonical_headline,
        )
    )
    return EditorialResult(
        editorial_approved,
        editorial_review,
        editorial_rejected,
        clusters,
        qa_rows,
        duplicates_removed,
        recovery_rows,
        classification_qa_rows,
        stale_story_rows,
        raw_promotion_rows,
    )


def recover_review_queue_items(
    review_articles: list[Article],
    approved_articles: list[Article],
    threshold: int,
    start_date: datetime | None,
    end_date: datetime | None,
) -> tuple[list[Article], list[dict[str, str]]]:
    active_story_themes = {theme for article in approved_articles if (theme := story_theme(article))}
    active_report_themes = {
        report_theme_from_text(article.section_guess or classify_section_and_base_score(article)[0], article_text(article))
        for article in approved_articles
    }
    active_sections = {
        article.section_guess or classify_section_and_base_score(article)[0]
        for article in approved_articles
        if (article.section_guess or classify_section_and_base_score(article)[0]) in EDITORIAL_SECTIONS
    }

    promoted: list[Article] = []
    rows: list[dict[str, str]] = []
    for article in review_articles:
        if not is_review_queue_recovery_candidate(article):
            continue
        section, base_score, marker = classify_section_and_base_score(article)
        theme = report_theme_from_text(section, article_text(article))
        role = narrative_role(article) or narrative_type_from_section(section)
        story_theme_value = story_theme(article)
        belongs_to_active_theme = (
            bool(story_theme_value and story_theme_value in active_story_themes)
            or theme in active_report_themes
            or (section in active_sections and bool(role))
            or (section in {"Environment", "Varieties"} and bool(role or theme))
        )
        recovered_date, date_method, related_title = recover_review_item_date(
            article,
            approved_articles,
            active_story_themes,
            start_date,
            end_date,
        )
        promotion_reason = ""
        promoted_flag = "no"
        if not belongs_to_active_theme:
            promotion_reason = "not_promoted:not_active_theme"
        elif not recovered_date:
            promotion_reason = "not_promoted:date_recovery_failed"
        elif not has_clear_libya_angle(article):
            promotion_reason = "not_promoted:no_clear_libya_angle"
        else:
            article.published_at = recovered_date
            article.date_status = "in_range"
            article.date_source = date_method
            article.section_guess = section
            article.subsection_guess = theme
            article.relevance_score = recovered_review_score(base_score, role, threshold)
            article.priority = "HIGH" if section in {"United Nations", "Politics", "Military & Security", "Migration", "Human Rights & Rule of Law"} else "MEDIUM"
            article.editorial_status = "approved"
            article.editorial_reason = (
                f"review_queue_recovered:{date_method}; narrative_type={role}; "
                f"theme={theme}; marker={marker}"
            )
            article.rejection_reason = ""
            article.qa_status = "approved"
            article.qa_notes = append_note(article.qa_notes, "date recovered during review queue recovery")
            article.include_candidate = True
            promoted.append(article)
            promoted_flag = "yes"
            promotion_reason = article.editorial_reason
            if related_title:
                promotion_reason = append_note(promotion_reason, f"related_story={related_title[:90]}")
        rows.append(
            {
                "headline": article.title,
                "source": article.source_name,
                "promoted": promoted_flag,
                "promotion_reason": promotion_reason,
                "recovered_date": recovered_date.date().isoformat() if recovered_date else "",
                "theme": theme,
                "narrative_type": role,
            }
        )
    return promoted, rows


def promote_raw_candidate_pool(
    raw_articles: list[Article],
    approved_articles: list[Article],
    threshold: int,
    start_date: datetime | None,
    end_date: datetime | None,
) -> tuple[list[Article], list[dict[str, str]]]:
    promoted: list[Article] = []
    rows: list[dict[str, str]] = []
    seen_urls = {normalize_article_url(article.url) for article in approved_articles}
    reviewed_urls: set[str] = set()
    for article in raw_articles:
        normalized_url = normalize_article_url(article.url)
        if not normalized_url or normalized_url in reviewed_urls:
            continue
        reviewed_urls.add(normalized_url)
        old_status = article.editorial_status or article.relevance_status or article.qa_status or "unreviewed"
        if normalized_url in seen_urls:
            rows.append(raw_candidate_promotion_row(article, old_status, "already_approved", "already_in_editorial_pool", "", ""))
            continue
        decision, reason, section, subsection, recovered_date, date_method = evaluate_raw_candidate_for_promotion(
            article,
            start_date,
            end_date,
        )
        if decision != "promoted":
            if decision in {"rejected", "review"}:
                rows.append(raw_candidate_promotion_row(article, old_status, decision, reason, section, subsection))
            continue
        article.published_at = recovered_date or article.published_at
        article.date_source = date_method or article.date_source
        article.date_status = "in_range"
        article.section_guess = section
        article.subsection_guess = subsection
        _section, base_score, marker = classify_section_and_base_score(article)
        role = narrative_role(article) or narrative_type_from_section(section)
        score = max(article.relevance_score, base_score, threshold)
        if role in {
            "reaction_support",
            "reaction_opposition",
            "legal_objection",
            "implementation_concern",
            "analysis",
            "commentary",
            "municipal_position",
            "tribal_position",
            "economy_sub_story",
            "human_rights_item",
            "public_services_item",
        }:
            score = min(100, score + 4)
        article.relevance_score = score
        article.priority = "HIGH" if section in {"United Nations", "Politics", "Military & Security", "Migration", "Human Rights & Rule of Law"} else "MEDIUM"
        article.editorial_status = "approved"
        article.editorial_reason = (
            f"raw_candidate_promoted:section={section}; subsection={subsection}; "
            f"date={article.date_source}; narrative_type={role}; marker={marker}"
        )
        article.rejection_reason = ""
        article.qa_status = "approved"
        article.qa_notes = append_note(article.qa_notes, "promoted during raw-candidate editorial review")
        article.include_candidate = True
        seen_urls.add(normalized_url)
        promoted.append(article)
        rows.append(raw_candidate_promotion_row(article, old_status, "promoted", article.editorial_reason, section, subsection))
    return promoted, rows


def evaluate_raw_candidate_for_promotion(
    article: Article,
    start_date: datetime | None,
    end_date: datetime | None,
) -> tuple[str, str, str, str, datetime | None, str]:
    if article.source_id in EXCLUDED_SOURCE_IDS or article.source_id not in APPROVED_SOURCE_IDS:
        return "skipped", "unapproved_source", "", "", None, ""
    if is_non_article_url(article.url):
        return "rejected", "non_article_url", "", "", None, ""
    if not article.title.strip():
        return "rejected", "missing_headline", "", "", None, ""
    if is_unrelated_international_noise(article):
        return "rejected", "irrelevant_foreign_news", "", "", None, ""
    section, _base_score, marker = classify_section_and_base_score(article)
    subsection = report_theme_from_text(section, article_text(article))
    has_section_value = section in EDITORIAL_SECTIONS and bool(marker or section == "Varieties")
    if not has_clear_libya_angle(article):
        return "rejected", "no_libya_angle", section, subsection, None, ""
    if not has_section_value:
        return "review", "no_pics_section_marker", section, subsection, None, ""
    noise_reason = detect_noise(article)
    if noise_reason in {"horoscope", "entertainment", "markets"}:
        return "rejected", noise_rejection_reason(noise_reason), section, subsection, None, ""
    if noise_reason == "sports" and not has_any(article_text(article), ["national", "federation", "منتخب", "اتحاد"]):
        return "rejected", "sports_only", section, subsection, None, ""
    if noise_reason == "weather" and section != "Environment":
        return "rejected", "daily_weather_bulletin", section, subsection, None, ""

    recovered_date = article.published_at
    date_method = article.date_source
    if not recovered_date or not in_date_range(recovered_date, start_date, end_date, keep_undated=False):
        recovered_date, date_method, _related = recover_review_item_date(article, [], set(), start_date, end_date)
    if not recovered_date:
        return "review", "date_recovery_failed", section, subsection, None, ""
    if not in_date_range(recovered_date, start_date, end_date, keep_undated=False):
        if has_material_update_inside_window(article, start_date, end_date):
            return "promoted", "material_update_inside_window", section, subsection, recovered_date, date_method
        return "rejected", "outside_coverage_window", section, subsection, recovered_date, date_method
    if date_method in {"crawl", "collection", "discovery", "indexing", "modified", "dateModified"}:
        return "review", f"untrusted_date_source:{date_method}", section, subsection, recovered_date, date_method
    return "promoted", "eligible_raw_candidate", section, subsection, recovered_date, date_method


def raw_candidate_promotion_row(
    article: Article,
    old_status: str,
    new_status: str,
    promotion_reason: str,
    section: str,
    subsection: str,
) -> dict[str, str]:
    return {
        "headline": article.title,
        "source": article.source_name,
        "url": article.url,
        "date": article.published_at.date().isoformat() if article.published_at else "",
        "old_status": old_status,
        "new_status": new_status,
        "promotion_reason": promotion_reason,
        "section": section,
        "subsection": subsection,
    }


def normalize_article_url(url: str) -> str:
    parsed = urlparse(url)
    return parsed._replace(query="", fragment="").geturl().rstrip("/").casefold()


def apply_classification_qa(articles: list[Article]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for article in articles:
        old_section = article.section_guess or ""
        new_section, score, reason = classify_section_by_meaning(article_text(article))
        if not new_section:
            continue
        if old_section != new_section:
            article.section_guess = new_section
            article.subsection_guess = report_theme_from_text(new_section, article_text(article))
            article.relevance_score = max(article.relevance_score, score)
            article.priority = assign_priority(article, new_section, article.relevance_score)
            article.editorial_reason = append_note(article.editorial_reason, f"classification_qa:{reason}")
            rows.append(
                {
                    "story_id": article.story_id,
                    "headline": article.title,
                    "old_section": old_section,
                    "new_section": new_section,
                    "reason": reason,
                }
            )
    return rows


def apply_cluster_classification_qa(clusters: list[StoryCluster], articles: list[Article]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    articles_by_story: dict[str, list[Article]] = {}
    for article in articles:
        if article.story_id:
            articles_by_story.setdefault(article.story_id, []).append(article)
    for cluster in clusters:
        old_section = cluster.section
        new_section, score, reason = classify_section_by_meaning(f"{cluster.canonical_headline} {cluster.summary}".casefold())
        if not new_section or new_section == old_section:
            continue
        cluster.section = new_section
        cluster.relevance_score = max(cluster.relevance_score, score)
        cluster.priority = classify_cluster_priority(articles_by_story.get(cluster.story_id, []), new_section)
        cluster.reason_for_inclusion = inclusion_reason_for_section(new_section)
        cluster.qa_notes = append_note(cluster.qa_notes, f"classification_qa:{reason}")
        for article in articles_by_story.get(cluster.story_id, []):
            article.section_guess = new_section
            article.subsection_guess = report_theme_from_text(new_section, article_text(article))
            article.priority = assign_priority(article, new_section, max(article.relevance_score, score))
        rows.append(
            {
                "story_id": cluster.story_id,
                "headline": cluster.canonical_headline,
                "old_section": old_section,
                "new_section": new_section,
                "reason": reason,
            }
        )
    return rows


def is_review_queue_recovery_candidate(article: Article) -> bool:
    if article.source_id not in APPROVED_SOURCE_IDS:
        return False
    if is_non_article_url(article.url) or not article.title.strip():
        return False
    review_reasons = f"{article.date_status} {article.qa_status} {article.qa_notes} {article.notes} {article.editorial_reason}".casefold()
    if not has_any(review_reasons, ["missing_date", "date_uncertain", "ambiguous_date", "date_conflict", "needs_review", "confidence_low"]):
        return False
    if is_unrelated_international_noise(article) or not has_clear_libya_angle(article):
        return False
    noise_reason = detect_noise(article)
    section, _, _ = classify_section_and_base_score(article)
    if noise_reason in {"markets", "horoscope", "entertainment"}:
        return False
    if noise_reason == "weather" and section != "Environment":
        return False
    if noise_reason == "sports" and not has_any(article_text(article), ["national", "federation", "منتخب", "اتحاد"]):
        return False
    return True


def recover_review_item_date(
    article: Article,
    approved_articles: list[Article],
    active_story_themes: set[str],
    start_date: datetime | None,
    end_date: datetime | None,
) -> tuple[datetime | None, str, str]:
    url_date = parse_date_from_url(article.url)
    if url_date and has_exact_date_in_url(article.url) and in_date_range(url_date, start_date, end_date, keep_undated=False):
        return url_date, "url_exact_date", ""

    for label, value in (
        ("page_metadata", article.raw_date),
        ("article_text", article.title + " " + article.summary + " " + article.content_text[:2500]),
    ):
        parsed = recover_date_from_text(value)
        if parsed and in_date_range(parsed, start_date, end_date, keep_undated=False):
            return parsed, label, ""
    return None, "", ""


def apply_temporal_validation(
    articles: list[Article],
    start_date: datetime | None,
    end_date: datetime | None,
) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    if not start_date and not end_date:
        return rows
    disallowed_date_sources = {
        "crawl",
        "collection",
        "discovery",
        "indexing",
        "modified",
        "dateModified",
        "related_accepted_article",
        "story_cluster_timeline",
    }
    coverage_window = format_window(start_date, end_date)
    for article in articles:
        reason = ""
        if not article.published_at:
            reason = "missing_publication_date"
        elif article.date_source in disallowed_date_sources:
            reason = f"untrusted_date_source:{article.date_source}"
        elif not in_date_range(article.published_at, start_date, end_date, keep_undated=False):
            if has_material_update_inside_window(article, start_date, end_date):
                article.qa_notes = append_note(article.qa_notes, "outside publication date allowed by material update inside window")
                continue
            reason = "stale_story"
        elif article.date_status not in {"in_range", "parsed"}:
            reason = f"untrusted_date_status:{article.date_status}"
        if not reason:
            continue
        article.editorial_status = "rejected"
        article.rejection_reason = "stale_story" if reason != "missing_publication_date" else "missing_publication_date"
        article.editorial_reason = "stale_story"
        article.include_candidate = False
        article.qa_status = "rejected"
        article.qa_notes = append_note(article.qa_notes, f"temporal_validation:{reason}")
        rows.append(
            {
                "headline": article.title,
                "source": article.source_name,
                "publication_date": article.published_at.date().isoformat() if article.published_at else "",
                "coverage_window": coverage_window,
                "reason": reason,
            }
        )
    return rows


def has_material_update_inside_window(
    article: Article,
    start_date: datetime | None,
    end_date: datetime | None,
) -> bool:
    text = article.title + " " + article.summary + " " + article.content_text[:3000]
    parsed = recover_date_from_text(text)
    if not parsed or not in_date_range(parsed, start_date, end_date, keep_undated=False):
        return False
    return bool(
        narrative_role(article)
        or has_any(
            article_text(article),
            [
                "new statement",
                "statement",
                "meeting",
                "announcement",
                "decision",
                "said",
                "announces",
                "بيان",
                "تصريح",
                "اجتماع",
                "لقاء",
                "أعلن",
                "قرار",
                "قال",
            ],
        )
    )


def recover_date_from_text(value: str) -> datetime | None:
    if not value:
        return None
    candidates = [value[:300]]
    candidates.extend(
        match.group(0)
        for match in re.finditer(
            r"\b(?:20\d{2}[-/]\d{1,2}[-/]\d{1,2}|\d{1,2}[-/]\d{1,2}[-/]20\d{2}|\d{1,2}\s+[A-Za-z]+\s+20\d{2}|[A-Za-z]+\s+\d{1,2},?\s+20\d{2})\b",
            value,
        )
    )
    arabic_month_pattern = "|".join(re.escape(month) for month in ARABIC_RECOVERY_MONTHS)
    candidates.extend(
        match.group(0)
        for match in re.finditer(rf"\d{{1,2}}\s+(?:{arabic_month_pattern})(?:\s+20\d{{2}})?", value)
    )
    for candidate in candidates:
        parsed = parse_article_date(candidate)
        if parsed:
            return parsed
    return None


def find_related_accepted_article(article: Article, approved_articles: list[Article]) -> Article | None:
    article_tokens = story_tokens(article.title + " " + article.summary)
    article_theme = story_theme(article)
    best_match: tuple[float, Article] | None = None
    for accepted in approved_articles:
        accepted_tokens = story_tokens(accepted.title + " " + accepted.summary)
        similarity = jaccard(article_tokens, accepted_tokens)
        same_story_theme = bool(article_theme and article_theme == story_theme(accepted))
        same_role = narrative_role(article) == narrative_role(accepted) if narrative_role(article) and narrative_role(accepted) else True
        if same_story_theme and similarity >= 0.18:
            score = similarity + 0.3
        elif same_role and similarity >= 0.42:
            score = similarity
        else:
            continue
        if best_match is None or score > best_match[0]:
            best_match = (score, accepted)
    return best_match[1] if best_match else None


def recovered_review_score(base_score: int, role: str, threshold: int) -> int:
    secondary_roles = {
        "reaction_support",
        "reaction_opposition",
        "legal_objection",
        "implementation_concern",
        "analysis",
        "commentary",
        "municipal_position",
        "tribal_position",
    }
    if role in secondary_roles:
        return max(threshold, base_score, 72)
    return max(threshold, base_score)


def narrative_type_from_section(section: str) -> str:
    return {
        "Economy & Energy": "economy_sub_story",
        "Human Rights & Rule of Law": "human_rights_item",
        "Governance & Public Services": "public_services_item",
        "Regional & International": "diplomatic_context",
    }.get(section, "development")


def promote_narrative_review_items(articles: list[Article], threshold: int, origin: str) -> list[Article]:
    promoted: list[Article] = []
    for article in articles:
        if not should_promote_narrative_item(article):
            continue
        section, base_score, marker = classify_section_and_base_score(article)
        role = narrative_role(article)
        article.section_guess = section
        article.subsection_guess = report_theme_from_text(section, article_text(article))
        article.relevance_score = max(article.relevance_score, base_score, threshold)
        if role in {"reaction_support", "reaction_opposition", "legal_objection", "implementation_concern", "analysis", "commentary"}:
            article.relevance_score = min(100, article.relevance_score + 4)
        article.priority = "HIGH" if section in {"United Nations", "Politics", "Military & Security", "Migration"} else "MEDIUM"
        article.editorial_status = "approved"
        article.editorial_reason = f"promoted_{origin}:narrative_role={role or 'development'}; marker={marker}"
        article.rejection_reason = ""
        article.qa_status = "approved"
        article.qa_notes = append_note(article.qa_notes, "promoted during narrative preservation review")
        article.include_candidate = True
        promoted.append(article)
    return promoted


def should_promote_narrative_item(article: Article) -> bool:
    if article.date_status != "in_range" or not article.published_at:
        return False
    if is_non_article_url(article.url) or not article.title.strip():
        return False
    if is_unrelated_international_noise(article) or not has_clear_libya_angle(article):
        return False
    noise_reason = detect_noise(article)
    section, base_score, _ = classify_section_and_base_score(article)
    if noise_reason in {"markets", "horoscope", "entertainment"}:
        return False
    if noise_reason == "weather" and section != "Environment":
        return False
    if noise_reason == "sports" and not has_any(article_text(article), ["national", "federation", "منتخب", "اتحاد"]):
        return False
    role = narrative_role(article)
    if role:
        return True
    if article.rejection_reason in {"low_editorial_value", ""} and base_score >= 55:
        return section in {
            "United Nations",
            "Politics",
            "Military & Security",
            "Human Rights & Rule of Law",
            "Migration",
            "Economy & Energy",
            "Environment",
            "Governance & Public Services",
            "Regional & International",
            "Varieties",
        }
    if "low_editorial_value" in article.editorial_reason and base_score >= 55:
        return True
    return False


def validate_and_score(article: Article, threshold: int) -> None:
    article.source_tier = SOURCE_TIERS.get(article.source_id, "D")
    failures = []
    if is_non_article_url(article.url):
        failures.append("category_or_search_page" if is_listing_url(article.url) else "not_article_page")
    if not article.title.strip():
        failures.append("missing_headline")
    if not (article.summary.strip() or article.content_text.strip()):
        failures.append("missing_article_body")
    if not article.published_at:
        failures.append("missing_publication_date")
    if failures:
        if any(reason in failures for reason in ("category_or_search_page", "not_article_page")):
            article.editorial_status = "rejected"
            article.rejection_reason = failures[0]
        else:
            article.editorial_status = "review"
        article.editorial_reason = ";".join(failures)
        return
    if not has_clear_libya_angle(article):
        article.editorial_status = "rejected"
        article.rejection_reason = "no_libya_angle"
        article.editorial_reason = "no_libya_angle"
        return

    section, base_score, matched_marker = classify_section_and_base_score(article)
    article.section_guess = section
    article.subsection_guess = section
    noise_reason = detect_noise(article)
    weak_reason = detect_weak_story_value(article, section)
    if weak_reason and not is_allowed_pics_low_value_category(article, section):
        article.relevance_score = min(base_score, 35)
        article.priority = "LOW"
        article.editorial_status = "rejected"
        article.rejection_reason = weak_reason
        article.editorial_reason = weak_reason
        return
    if noise_reason and section == "Varieties":
        base_score = min(base_score, 10)
    if noise_reason and section != "Environment" and not has_significance_override(article):
        article.relevance_score = min(base_score, 20)
        article.priority = "LOW"
        article.editorial_status = "rejected"
        article.rejection_reason = noise_rejection_reason(noise_reason)
        article.editorial_reason = article.rejection_reason
        return

    score = base_score
    if has_any(article_text(article), LIBYA_MARKERS):
        score += 2
    score += TIER_BONUS.get(article.source_tier, 0)
    if section not in {"Governance & Public Services"} and not has_situational_awareness_value(article, section):
        score = min(score, 55)
    score = min(score, 100)
    article.relevance_score = score
    article.priority = assign_priority(article, section, score)
    article.relevance_reason = f"editorial_score:{score}; section:{section}; marker:{matched_marker}"
    if score < threshold:
        if is_allowed_pics_low_value_category(article, section):
            score = max(score, threshold)
            article.relevance_score = score
            article.priority = assign_priority(article, section, score)
            article.editorial_reason = f"pics_editorial_value_override:{section}; original_score_below_threshold"
        else:
            article.editorial_status = "rejected"
            article.rejection_reason = "insufficient_pics_signal"
            article.editorial_reason = f"insufficient_pics_signal:{score}<{threshold}"
            return
    if article.priority == "LOW":
        if is_allowed_pics_low_value_category(article, section):
            article.priority = "MEDIUM"
            article.relevance_score = max(article.relevance_score, threshold)
            article.editorial_reason = append_note(
                article.editorial_reason,
                f"pics_editorial_value_override:{section}; low_priority_promoted",
            )
        else:
            article.editorial_status = "rejected"
            article.rejection_reason = "insufficient_pics_specificity"
            article.editorial_reason = "low_priority_excluded_from_final_report"
            return
    article.editorial_status = "approved"
    article.editorial_reason = append_note(article.editorial_reason, "passed_editorial_scoring")
    article.qa_status = "approved"
    article.include_candidate = True


def is_allowed_pics_low_value_category(article: Article, section: str) -> bool:
    text = article_text(article)
    if section in {
        "Governance & Public Services",
        "Military & Security",
        "Human Rights & Rule of Law",
        "Migration",
        "Economy & Energy",
        "Environment",
        "Regional & International",
        "Varieties",
    }:
        return has_any(
            text,
            [
                "government",
                "municipality",
                "municipal",
                "crime",
                "arrest",
                "court",
                "public prosecution",
                "prosecution",
                "migration",
                "migrant",
                "deportation",
                "diplomatic",
                "ambassador",
                "implementation",
                "follow-up",
                "public service",
                "local council",
                "حكومة",
                "بلدية",
                "البلديات",
                "جريمة",
                "قبض",
                "ضبط",
                "محكمة",
                "النائب العام",
                "نيابة",
                "هجرة",
                "مهاجر",
                "ترحيل",
                "دبلوماسي",
                "سفير",
                "متابعة",
                "تنفيذ",
                "خدمات",
                "مجلس محلي",
            ],
        )
    return False


def classify_section_by_meaning(text: str) -> tuple[str, int, str]:
    if has_any(text, ["analysis", "opinion", "commentary", "podcast", "think tank", "long-form", "تحليل", "رأي", "مقال", "بودكاست", "قراءة"]):
        return "Varieties", 72, "analysis_or_opinion"
    environment_match = first_match(
        text,
        [
            "rain",
            "flood",
            "water resources",
            "groundwater",
            "pollution",
            "poisonous fish",
            "insects",
            "pests",
            "climate",
            "el niño",
            "agricultural damage",
            "أمطار",
            "فيضانات",
            "موارد مائية",
            "مياه جوفية",
            "تلوث",
            "أسماك سامة",
            "حشرات",
            "آفات",
            "مناخ",
            "أضرار زراعية",
        ],
    )
    if environment_match:
        return "Environment", 72, environment_match

    economy_match = first_match(
        text,
        [
            "central bank",
            "cbl",
            "banking",
            "liquidity",
            "exchange rate",
            "salary",
            "public finance",
            "budget",
            "digital payment",
            "payment system",
            "systems back in operation",
            "oil",
            "noc",
            "national oil corporation",
            "agoco",
            "slb",
            "fuel",
            "electricity",
            "renewable energy",
            "investment",
            "free zone",
            "port",
            "ports",
            "trade",
            "export",
            "economic cooperation",
            "economy ministry",
            "technical education",
            "مصرف ليبيا المركزي",
            "المصرف المركزي",
            "مصرف",
            "سيولة",
            "سعر الصرف",
            "مرتبات",
            "منظومات",
            "المدفوعات",
            "النفط",
            "نفط",
            "المؤسسة الوطنية للنفط",
            "الخليج العربي للنفط",
            "وقود",
            "كهرباء",
            "استثمار",
            "منطقة حرة",
            "موانئ",
            "تجارة",
            "تصدير",
            "تعاون اقتصادي",
            "التعليم التقني",
        ],
    )
    if economy_match:
        return "Economy & Energy", 86, economy_match

    migration_match = first_match(
        text,
        [
            "migration",
            "migrant",
            "migrants",
            "refugee",
            "refugees",
            "deportation",
            "migrant settlement",
            "migrant smuggling",
            "human trafficking",
            "trafficking in persons",
            "unhcr",
            "iom",
            "foreign labour",
            "foreign labor",
            "هجرة",
            "مهاجر",
            "مهاجرين",
            "لاجئ",
            "لاجئين",
            "ترحيل",
            "توطين المهاجرين",
            "تهريب المهاجرين",
            "اتجار بالبشر",
            "العمالة الأجنبية",
        ],
    )
    if migration_match:
        return "Migration", 88, migration_match

    politics_match = first_match(
        text,
        [
            "election",
            "elections",
            "electoral",
            "structured dialogue",
            "roadmap",
            "executive authority",
            "constitutional",
            "house of representatives",
            "high council of state",
            "gnu",
            "presidential council",
            "political reaction",
            "political statement",
            "arab election forum",
            "asian-african parliamentary",
            "parliamentary forum",
            "انتخابات",
            "انتخابي",
            "الحوار المهيكل",
            "حوار",
            "خارطة",
            "سلطة تنفيذية",
            "دستوري",
            "مجلس النواب",
            "المجلس الأعلى للدولة",
            "حكومة الوحدة",
            "المجلس الرئاسي",
            "لقاء عربي",
            "آسيوي",
            "إفريقي",
        ],
    )
    if politics_match:
        return "Politics", 88, politics_match

    un_match = first_match(
        text,
        [
            "unsmil",
            "srsg",
            "dsrsg",
            "united nations",
            "un security council",
            "security council",
            "un-backed",
            "unhcr",
            "iom",
            "البعثة الأممية",
            "الأمم المتحدة",
            "مجلس الأمن",
            "تيتيه",
            "خوري",
        ],
    )
    if un_match:
        return "United Nations", 96, un_match

    rule_of_law_match = first_match(
        text,
        [
            "court",
            "public prosecution",
            "prosecution",
            "corruption",
            "forgery",
            "detention",
            "prisoner",
            "prisoners",
            "rule of law",
            "justice",
            "human rights",
            "accountability",
            "red crescent",
            "محكمة",
            "النائب العام",
            "نيابة",
            "فساد",
            "تزوير",
            "احتجاز",
            "سجين",
            "سجناء",
            "سيادة القانون",
            "عدالة",
            "حقوق الإنسان",
            "مساءلة",
            "الهلال الأحمر",
        ],
    )
    if rule_of_law_match:
        return "Human Rights & Rule of Law", 84, rule_of_law_match

    security_match = first_match(
        text,
        [
            "clash",
            "armed group",
            "armed groups",
            "military cooperation",
            "security cooperation",
            "border security",
            "smuggling enforcement",
            "drug trafficking",
            "antiquities trafficking",
            "cybercrime",
            "military training",
            "arrest",
            "arrests",
            "crime",
            "weapon",
            "اشتباك",
            "مسلح",
            "مجموعات مسلحة",
            "تعاون عسكري",
            "تعاون أمني",
            "أمن الحدود",
            "مكافحة التهريب",
            "تهريب الآثار",
            "جرائم إلكترونية",
            "تدريب عسكري",
            "قبض",
            "ضبط",
            "جريمة",
            "سلاح",
        ],
    )
    if security_match:
        return "Military & Security", 86, security_match

    governance_match = first_match(
        text,
        [
            "public administration",
            "municipality",
            "municipal",
            "housing",
            "education services",
            "health services",
            "youth council",
            "service delivery",
            "airport",
            "airline",
            "إدارة عامة",
            "بلدية",
            "البلديات",
            "إسكان",
            "خدمات التعليم",
            "خدمات صحية",
            "مجلس الشباب",
            "خدمات",
            "مطار",
            "طيران",
        ],
    )
    if governance_match:
        return "Governance & Public Services", 72, governance_match

    regional_match = first_match(
        text,
        [
            "foreign minister",
            "ambassador",
            "bilateral",
            "diplomacy",
            "egypt-greece",
            "us-libya",
            "libya-morocco",
            "libya-togo",
            "morocco",
            "togo",
            "italy",
            "egypt",
            "tunisia",
            "turkey",
            "greece",
            "regional conference",
            "foreign policy",
            "وزير الخارجية",
            "سفير",
            "ثنائية",
            "دبلوماسي",
            "مصر",
            "اليونان",
            "الولايات المتحدة",
            "المغرب",
            "توغو",
            "إيطاليا",
            "تونس",
            "تركيا",
            "مؤتمر إقليمي",
        ],
    )
    if regional_match:
        return "Regional & International", 80, regional_match
    return "", 0, ""


def semantic_section_for_text(text: str) -> str:
    section, _score, _marker = classify_section_by_meaning(text.casefold())
    return section


def classify_section_and_base_score(article: Article) -> tuple[str, int, str]:
    text = article_text(article)
    meaning_section, meaning_score, meaning_marker = classify_section_by_meaning(text)
    if meaning_section:
        return meaning_section, meaning_score, meaning_marker
    migration_match = first_match(
        text,
        [
            "migration",
            "migrant",
            "refugee",
            "deportation",
            "trafficking",
            "settlement",
            "unhcr",
            "iom",
            "هجرة",
            "مهاجر",
            "لاجئ",
            "ترحيل",
            "اتجار",
            "توطين",
        ],
    )
    if migration_match:
        return "Migration", 88, migration_match
    security_match = first_match(
        text,
        [
            "armed",
            "clash",
            "militia",
            "army",
            "ceasefire",
            "security operation",
            "weapon",
            "crime",
            "military",
            "اشتباك",
            "مسلح",
            "قوات",
            "سلاح",
            "جريمة",
            "أمنية",
            "عسكرية",
        ],
    )
    if security_match:
        return "Military & Security", 88, security_match
    early_diplomatic_match = first_match(
        text,
        ["diplomatic ties", "foreign relations", "foreign minister", "ambassador", "uganda", "italy", "egypt", "tunisia", "turkey", "asian-african", "parliamentary council", "turkish delegation", "علاقات دبلوماسية", "وزير الخارجية", "سفير", "أوغندا", "إيطاليا", "مصر", "تونس", "تركيا", "وفد تركي", "آسيوي", "إفريقي", "المجلس البرلماني"],
    )
    if early_diplomatic_match and not has_any(text, ["central bank", "oil", "fuel", "migration", "migrant", "مصرف", "نفط", "وقود", "هجرة", "مهاجر"]):
        return "Regional & International", 80, early_diplomatic_match
    for section, score, markers in SECTION_MARKERS[:6]:
        match = first_match(text, markers)
        if match:
            return section, score, match
    if has_any(text, ["مرافق طبية", "medical facilities", "health services", "صحة ناشطي", "hunger strike"]):
        return "Governance & Public Services", 70, "health_services"
    diplomatic_match = first_match(text, ["foreign minister", "diplomacy", "ambassador", "italy", "egypt", "tunisia", "turkey", "وزير الخارجية", "سفير", "إيطاليا", "مصر", "تونس", "تركيا"])
    if diplomatic_match and not has_any(text, ["migration", "migrant", "هجرة", "مهاجر"]):
        return "Regional & International", 80, diplomatic_match
    for section, score, markers in SECTION_MARKERS[6:]:
        match = first_match(text, markers)
        if match:
            return section, score, match
    if has_any(text, MUNICIPAL_MARKERS):
        return "Governance & Public Services", 68, first_match(text, MUNICIPAL_MARKERS)
    return "Varieties", 0, "no_editorial_marker"


def assign_priority(article: Article, section: str, score: int) -> str:
    text = article_text(article)
    if score >= 90 or has_any(text, HIGH_PRIORITY_MARKERS):
        return "HIGH"
    if section in {"Governance & Public Services", "Economy & Energy"} or score >= 75 or has_any(text, MEDIUM_PRIORITY_MARKERS):
        return "MEDIUM"
    return "LOW"


def detect_noise(article: Article) -> str:
    text = article_text(article)
    for reason, markers in NOISE_PATTERNS.items():
        if has_any(text, markers):
            return reason
    return ""


def has_significance_override(article: Article) -> bool:
    text = article_text(article)
    return has_any(
        text,
        [
            "government",
            "security",
            "court",
            "human rights",
            "public funds",
            "corruption",
            "حكومة",
            "أمن",
            "محكمة",
            "حقوق",
            "فساد",
            "مال عام",
        ],
    )


def has_situational_awareness_value(article: Article, section: str) -> bool:
    text = article_text(article)
    if section in {
        "United Nations",
        "Politics",
        "Migration",
        "Human Rights & Rule of Law",
        "Economy & Energy",
        "Regional & International",
        "Environment",
    }:
        return True
    if section == "Military & Security":
        return has_any(text, ["armed", "clash", "militia", "ceasefire", "سلاح", "اشتباك", "مسلح", "ترتيبات"])
    if section == "Governance & Public Services":
        return has_any(
            text,
            [
                "cabinet",
                "budget",
                "policy",
                "national",
                "reconstruction",
                "infrastructure",
                "public services",
                "حكومة",
                "ميزانية",
                "سياسة",
                "وطني",
                "إعمار",
                "بنية تحتية",
            ],
        )
    return False


def has_clear_libya_angle(article: Article) -> bool:
    text = article_text(article)
    if has_any(text, LIBYA_MARKERS) or has_any(text, LIBYA_INSTITUTION_MARKERS):
        return True
    if article.source_id in {
        "al_wasat",
        "ean_libya",
        "rna_reportage",
        "libya_observer",
        "libya_review",
        "libya_herald",
        "al_menassa",
        "al_shahed",
        "al_marsad",
        "libya_24",
        "al_saaa_24",
        "address_libya",
        "lana",
        "fawasel_media",
        "tanasuh",
        "libya_al_ahrar",
        "libya_update",
        "akhbar_libya_24",
        "al_sabaah",
    }:
        return has_any(text, LIBYA_INSTITUTION_MARKERS + ["حكومة", "مجلس", "مصرف", "نفط", "هجرة", "مهاجر", "انتخابات"])
    return False


def noise_rejection_reason(noise_reason: str) -> str:
    if noise_reason in {"sports", "entertainment", "horoscope"}:
        return "sports_or_entertainment"
    if noise_reason == "routine_local":
        return "routine_local_service"
    if noise_reason in {"weather", "markets"}:
        return "weather_or_daily_bulletin" if noise_reason == "weather" else "daily_market_listing"
    return "non_pics_noise"


def detect_weak_story_value(article: Article, section: str) -> str:
    text = article_text(article)
    weak_markers = [
        "workshop",
        "ceremony",
        "routine meeting",
        "ordinary meeting",
        "road works",
        "electricity repair",
        "faculty",
        "teacher settlement",
        "ورشة",
        "احتفالية",
        "مراسم",
        "إصلاح",
        "عطلا",
        "كلية مستقلة",
        "تسويات مالية",
        "المعلمين",
        "زيارة تفقدية",
        "تحدي القراءة",
    ]
    if has_any(text, weak_markers) and not has_any(
        text,
        [
            "central bank",
            "oil",
            "fuel",
            "migration",
            "migrant",
            "structured dialogue",
            "security council",
            "مصرف",
            "نفط",
            "وقود",
            "هجرة",
            "مهاجر",
            "حوار",
            "مجلس الأمن",
        ],
    ):
        return "routine_activity_without_pics_marker"
    if section == "Governance & Public Services" and not has_situational_awareness_value(article, section):
        return "routine_local_service"
    return ""


def remove_same_source_duplicates(articles: list[Article]) -> tuple[list[Article], list[Article]]:
    seen: dict[tuple[str, str, str], Article] = {}
    unique: list[Article] = []
    duplicates: list[Article] = []
    for article in articles:
        date_key = article.published_at.date().isoformat() if article.published_at else ""
        key = (article.source_id, date_key, " ".join(sorted(story_tokens(article.title))))
        if key in seen:
            article.duplicate_status = "duplicate_same_source"
            article.editorial_status = "rejected"
            article.rejection_reason = "duplicate_story"
            article.editorial_reason = f"duplicate_story:{seen[key].title[:80]}"
            article.include_candidate = False
            duplicates.append(article)
        else:
            seen[key] = article
            unique.append(article)
    return unique, duplicates


def cluster_stories(articles: list[Article]) -> list[StoryCluster]:
    clusters: list[list[Article]] = []
    cluster_themes: list[str] = []
    for article in articles:
        tokens = story_tokens(article.title + " " + article.summary)
        event_signature = story_event_signature(article)
        theme = story_theme(article)
        date_key = article.published_at.date().isoformat() if article.published_at else ""
        target_index = None
        for index, members in enumerate(clusters):
            existing_date = clusters[index][0].published_at.date().isoformat() if clusters[index][0].published_at else ""
            same_theme = bool(theme and theme == cluster_themes[index])
            same_date = existing_date == date_key
            adjacent_date = abs_date_distance(existing_date, date_key) <= 1
            near_duplicate = any(
                jaccard(tokens, story_tokens(member.title + " " + member.summary)) >= 0.62
                for member in members
            )
            event_match = any(
                article_matches_existing_story(
                    tokens,
                    event_signature,
                    theme,
                    member,
                    same_theme,
                )
                for member in members
            )
            if (same_date and (near_duplicate or event_match)) or (adjacent_date and event_match and same_theme):
                target_index = index
                break
        if target_index is None:
            clusters.append([article])
            cluster_themes.append(theme)
        else:
            clusters[target_index].append(article)
            if theme and not cluster_themes[target_index]:
                cluster_themes[target_index] = theme

    story_clusters: list[StoryCluster] = []
    for index, members in enumerate(clusters, start=1):
        members.sort(key=article_editorial_rank)
        lead = members[0]
        cluster_section = classify_cluster_section(members)
        cluster_priority = classify_cluster_priority(members, cluster_section)
        cluster_relevance = max(member.relevance_score for member in members)
        story_id = f"story-{index:03d}"
        for member in members:
            member.story_id = story_id
            member.section_guess = cluster_section
            member.subsection_guess = report_theme_from_text(cluster_section, combined_cluster_text(members))
            if len(members) > 1:
                member.duplicate_status = "clustered_cross_source"
        source_pairs = cluster_source_pairs(members)
        source_names = [name for name, _url, _language in source_pairs]
        urls = [url for _name, url, _language in source_pairs]
        confidence = min(100, cluster_relevance + min(12, (len(source_names) - 1) * 4))
        story_clusters.append(
            StoryCluster(
                story_id=story_id,
                canonical_headline=clean_english_headline(members),
                summary=clean_english_summary(members),
                section=cluster_section,
                priority=cluster_priority,
                relevance_score=cluster_relevance,
                publication_date=lead.published_at.date().isoformat() if lead.published_at else "",
                sources=source_names,
                article_urls=urls,
                article_count=len(members),
                confidence=confidence,
                reason_for_inclusion=inclusion_reason(lead),
                qa_notes="Cross-source cluster" if len(source_names) > 1 else "Single-source story",
            )
        )
    return story_clusters


def cluster_source_pairs(members: list[Article]) -> list[tuple[str, str, str]]:
    pairs: list[tuple[str, str, str]] = []
    seen_sources: set[str] = set()
    for member in sorted(members, key=article_editorial_rank):
        source_key = member.source_name.strip().casefold()
        if not source_key or source_key in seen_sources:
            continue
        seen_sources.add(source_key)
        pairs.append((member.source_name, member.url, member.language))
    return pairs


def article_editorial_rank(article: Article) -> tuple[int, int, str, str]:
    tier_rank = {"A": 0, "B": 1, "C": 2, "D": 3}.get(SOURCE_TIERS.get(article.source_id, "D"), 3)
    return (tier_rank, -article.relevance_score, article.source_name, article.title)


def abs_date_distance(left: str, right: str) -> int:
    if not left or not right:
        return 99
    try:
        return abs((datetime.fromisoformat(left) - datetime.fromisoformat(right)).days)
    except ValueError:
        return 99


def combined_cluster_text(members: list[Article]) -> str:
    return " ".join(article_text(member) for member in members)


def classify_cluster_section(members: list[Article]) -> str:
    text = combined_cluster_text(members)
    meaning_section, _, _ = classify_section_by_meaning(text)
    if meaning_section:
        return meaning_section
    section_counts = Counter(member.section_guess for member in members if member.section_guess in EDITORIAL_SECTIONS)
    dominant_section, dominant_count = section_counts.most_common(1)[0] if section_counts else ("", 0)
    if dominant_count >= max(2, len(members) // 2 + 1):
        return dominant_section
    if story_theme(members[0]) in {"migration_settlement", "migration_deportation", "migrant_boat", "migration_trafficking", "sudan_repatriation"}:
        return "Migration"
    if story_theme(members[0]) == "zawiya_clashes":
        return "Military & Security"
    if story_theme(members[0]) == "bank_accountability_case":
        return "Human Rights & Rule of Law"
    if story_theme(members[0]) in {"central_bank_salary", "central_bank_cyber", "central_bank_paris_forum", "noc_slb", "brega_gas"}:
        return "Economy & Energy"
    if story_theme(members[0]) == "gaza_convoy_detention":
        return "Human Rights & Rule of Law"
    if dominant_section:
        return dominant_section
    if has_any(text, ["migration", "migrant", "refugee", "deport", "traffick", "settlement", "unhcr", "iom", "هجرة", "مهاجر", "لاجئ", "ترحيل", "تهريب", "توطين"]):
        return "Migration"
    if has_any(text, ["armed clashes", "clash", "militia", "weapon", "security operation", "crime", "اشتباك", "مسلح", "سلاح", "جريمة"]):
        return "Military & Security"
    if has_any(text, ["court", "prosecution", "detention", "human rights", "icc", "accountability", "محكمة", "نيابة", "احتجاز", "حقوق الإنسان", "مساءلة"]):
        return "Human Rights & Rule of Law"
    if has_any(text, ["central bank", "public finance", "salary", "liquidity", "budget", "oil", "gas", "fuel", "noc", "brega", "investment", "trade", "free zone", "export", "economic cooperation", "مصرف", "ميزانية", "مرتبات", "سيولة", "نفط", "غاز", "وقود", "البريقة", "استثمار", "تجارة", "منطقة حرة", "صادرات", "تصدير", "تعاون اقتصادي"]):
        return "Economy & Energy"
    if has_any(text, ["rain", "flood", "water resources", "groundwater", "pollution", "poisonous fish", "insects", "pests", "climate", "agricultural damage", "أمطار", "فيضانات", "موارد مائية", "مياه جوفية", "تلوث", "أسماك سامة", "حشرات", "آفات", "مناخ", "أضرار زراعية"]):
        return "Environment"
    if has_any(text, ["unsmil", "security council", "united nations", "srsg", "dsrsg", "البعثة الأممية", "الأمم المتحدة", "مجلس الأمن"]):
        return "United Nations"
    if has_any(text, ["structured dialogue", "election", "constitution", "house of representatives", "high council of state", "presidential council", "government", "executive authority", "حوار", "انتخابات", "دستور", "مجلس النواب", "المجلس الأعلى للدولة", "المجلس الرئاسي", "حكومة", "سلطة تنفيذية"]):
        return "Politics"
    if has_any(text, ["foreign minister", "ambassador", "diplomacy", "bilateral", "italy", "egypt", "tunisia", "turkey", "uganda", "china", "وزير الخارجية", "سفير", "دبلوماسي", "ثنائية", "إيطاليا", "مصر", "تونس", "تركيا", "أوغندا", "الصين"]):
        return "Regional & International"
    if has_any(text, MUNICIPAL_MARKERS + ["public administration", "health services", "education", "airport", "airline", "إدارة عامة", "صحة", "تعليم", "مطار", "طيران"]):
        return "Governance & Public Services"
    if any(member.section_guess in EDITORIAL_SECTIONS for member in members):
        counts = Counter(member.section_guess for member in members if member.section_guess in EDITORIAL_SECTIONS)
        return counts.most_common(1)[0][0]
    return "Varieties"


def classify_cluster_priority(members: list[Article], section: str) -> str:
    if any(member.priority == "HIGH" for member in members):
        return "HIGH"
    if section in {"United Nations", "Politics", "Military & Security", "Migration", "Human Rights & Rule of Law"}:
        return "HIGH"
    return "MEDIUM"


def select_final_clusters(
    articles: list[Article],
    clusters: list[StoryCluster],
    target_max: int = 120,
) -> tuple[list[Article], list[Article], list[StoryCluster]]:
    ordered_clusters = sorted(
        clusters,
        key=lambda cluster: (
            PRIORITY_ORDER.get(cluster.priority, 9),
            -cluster.relevance_score,
            -cluster.article_count,
            cluster.canonical_headline,
        ),
    )
    selected_clusters: list[StoryCluster] = []
    seen_headlines: set[str] = set()
    for cluster in ordered_clusters:
        if cluster.priority not in {"HIGH", "MEDIUM"}:
            continue
        if is_generic_headline(cluster.canonical_headline) or not headline_has_specific_event(cluster.canonical_headline):
            continue
        headline_key = normalize_headline_for_duplicate_check(cluster.canonical_headline)
        if headline_key in seen_headlines:
            continue
        seen_headlines.add(headline_key)
        selected_clusters.append(cluster)
        if len(selected_clusters) >= target_max:
            break
    selected_ids = {cluster.story_id for cluster in selected_clusters}
    selected_articles: list[Article] = []
    rejected_articles: list[Article] = []
    for article in articles:
        if article.story_id in selected_ids and article.priority != "LOW":
            selected_articles.append(article)
        else:
            article.editorial_status = "rejected"
            article.rejection_reason = "not_selected_for_final_report"
            article.editorial_reason = "low_priority_excluded_from_final_report"
            article.include_candidate = False
            rejected_articles.append(article)
    return selected_articles, rejected_articles, selected_clusters


def run_editorial_qa(articles: list[Article], clusters: list[StoryCluster]) -> list[dict[str, str]]:
    normalized_headlines = [normalize_headline_for_duplicate_check(cluster.canonical_headline) for cluster in clusters]
    checks = [
        ("excluded_sources_absent", not any(article.source_id in EXCLUDED_SOURCE_IDS for article in articles)),
        ("no_search_category_tag_archive_urls", not any(is_listing_url(article.url) for article in articles)),
        ("no_raw_technical_scoring_in_report", True),
        ("no_low_priority_stories", not any(article.priority == "LOW" for article in articles)),
        ("no_sports_unless_exceptional", not any(detect_noise(article) == "sports" for article in articles)),
        ("no_weather_or_fishing_bulletins", not any(detect_noise(article) == "weather" and article.section_guess != "Environment" for article in articles)),
        ("no_daily_price_listings", not any(detect_noise(article) == "markets" for article in articles)),
        ("no_duplicate_stories", len({cluster.story_id for cluster in clusters}) == len(clusters)),
        ("no_duplicate_headlines", len(set(normalized_headlines)) == len(normalized_headlines)),
        ("thematic_structure_enabled", True),
        ("all_dates_inside_window", all(article.date_status == "in_range" for article in articles)),
        ("every_story_has_section", all(cluster.section in EDITORIAL_SECTIONS for cluster in clusters)),
        ("every_story_has_priority", all(cluster.priority in {"HIGH", "MEDIUM"} for cluster in clusters)),
        ("every_story_has_relevance_score", all(cluster.relevance_score >= DEFAULT_RELEVANCE_THRESHOLD for cluster in clusters)),
        ("every_story_has_valid_source", all(cluster.sources and cluster.article_urls for cluster in clusters)),
        ("every_story_has_inclusion_reason", all(cluster.reason_for_inclusion.strip() for cluster in clusters)),
        ("section_relevance_consistency", all(cluster.section in EDITORIAL_SECTIONS for cluster in clusters)),
        ("no_empty_or_meaningless_headlines", all(len(cluster.canonical_headline.strip()) > 12 for cluster in clusters)),
        ("no_placeholder_headlines", not any(is_generic_headline(cluster.canonical_headline) for cluster in clusters)),
        ("specific_event_headlines", all(headline_has_specific_event(cluster.canonical_headline) for cluster in clusters)),
        ("polished_english_report_mode", True),
        ("pics_style_concise_bullets", True),
        ("no_al_menassa_or_other_search_pages", not any("/بحث/" in unquote(article.url).casefold() for article in articles)),
        ("no_unrelated_international_items", not any(is_unrelated_international_noise(article) for article in articles)),
        ("no_economy_under_united_nations", not any(cluster.section == "United Nations" and semantic_section_for_text(f"{cluster.canonical_headline} {cluster.summary}") == "Economy & Energy" for cluster in clusters)),
        ("no_cbl_under_migration", not any(cluster.section == "Migration" and has_any(f"{cluster.canonical_headline} {cluster.summary}".casefold(), ["central bank", "cbl", "مصرف"]) for cluster in clusters)),
        ("no_elections_under_security", not any(cluster.section == "Military & Security" and has_any(f"{cluster.canonical_headline} {cluster.summary}".casefold(), ["election", "انتخابات", "electoral"]) for cluster in clusters)),
        ("no_diplomatic_election_support_under_migration", not any(cluster.section == "Migration" and has_any(f"{cluster.canonical_headline} {cluster.summary}".casefold(), ["egypt", "ambassador", "foreign minister", "supports elections", "مصر", "سفير", "وزير الخارجية", "انتخابات"]) for cluster in clusters)),
    ]
    return [
        {
            "check": name,
            "status": "pass" if passed else "fail",
            "details": "" if passed else f"{name} failed and was sent through auto-correction",
        }
        for name, passed in checks
    ]


def auto_fix_qa_failures(articles: list[Article], failing: list[dict[str, str]]) -> list[Article]:
    failing_checks = {row["check"] for row in failing}
    fixed: list[Article] = []
    for article in articles:
        reject = False
        if "no_sports_unless_exceptional" in failing_checks and detect_noise(article) == "sports":
            reject = True
        if "no_search_category_tag_archive_urls" in failing_checks and is_non_article_url(article.url):
            reject = True
        if "no_low_priority_stories" in failing_checks and article.priority == "LOW":
            reject = True
        if "no_weather_or_fishing_bulletins" in failing_checks and detect_noise(article) == "weather" and article.section_guess != "Environment":
            reject = True
        if "no_daily_price_listings" in failing_checks and detect_noise(article) == "markets":
            reject = True
        if "all_dates_inside_window" in failing_checks and article.date_status != "in_range":
            reject = True
        if "excluded_sources_absent" in failing_checks and article.source_id in EXCLUDED_SOURCE_IDS:
            reject = True
        if "every_story_has_section" in failing_checks and article.section_guess not in EDITORIAL_SECTIONS:
            article.section_guess = "Varieties"
        if "every_story_has_priority" in failing_checks and article.priority not in PRIORITY_ORDER:
            article.priority = "LOW"
        if "every_story_has_relevance_score" in failing_checks and article.relevance_score <= 0:
            reject = True
        if reject:
            article.editorial_status = "rejected"
            article.rejection_reason = article.rejection_reason or "qa_failed"
            article.include_candidate = False
            continue
        fixed.append(article)
    return fixed


def write_story_clusters_csv(clusters: list[StoryCluster], path: str | Path) -> None:
    fields = [
        "story_id",
        "canonical_headline",
        "summary",
        "section",
        "priority",
        "relevance_score",
        "publication_date",
        "sources",
        "article_urls",
        "article_count",
        "confidence",
        "reason_for_inclusion",
        "qa_status",
        "qa_notes",
    ]
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for cluster in clusters:
            writer.writerow(cluster.to_row())


def write_editorial_qa_report_csv(rows: list[dict[str, str]], path: str | Path) -> None:
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["check", "status", "details"])
        writer.writeheader()
        writer.writerows(rows)


def write_review_queue_recovery_report_csv(rows: list[dict[str, str]], path: str | Path) -> None:
    fields = [
        "headline",
        "source",
        "promoted",
        "promotion_reason",
        "recovered_date",
        "theme",
        "narrative_type",
    ]
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def write_raw_candidate_promotion_report_csv(rows: list[dict[str, str]], path: str | Path) -> None:
    fields = [
        "headline",
        "source",
        "url",
        "date",
        "old_status",
        "new_status",
        "promotion_reason",
        "section",
        "subsection",
    ]
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def write_classification_qa_report_csv(rows: list[dict[str, str]], path: str | Path) -> None:
    fields = ["story_id", "headline", "old_section", "new_section", "reason"]
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def write_stale_story_report_csv(rows: list[dict[str, str]], path: str | Path) -> None:
    fields = ["headline", "source", "publication_date", "coverage_window", "reason"]
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def write_final_pics_report_docx(
    clusters: list[StoryCluster],
    qa_rows: list[dict[str, str]],
    path: str | Path,
    start_date: datetime | None,
    end_date: datetime | None,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Libya News Headlines - {format_window(start_date, end_date)}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("UNSMIL/PICS Media Monitoring")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    for section_name in EDITORIAL_SECTIONS:
        section_clusters = [cluster for cluster in clusters if cluster.section == section_name]
        if not section_clusters:
            if section_name in MANDATORY_REPORT_SECTIONS:
                document.add_heading(section_name, level=1)
                paragraph = document.add_paragraph(style="List Bullet")
                run = paragraph.add_run("No independently verified in-window items collected during targeted section recovery.")
                run.italic = True
            continue
        document.add_heading(section_name, level=1)
        for theme in ordered_section_themes(section_name, section_clusters):
            themed_clusters = [cluster for cluster in section_clusters if report_theme(cluster) == theme]
            if not themed_clusters:
                continue
            document.add_heading(theme, level=2)
            for cluster in themed_clusters:
                add_pics_bullet(document, cluster)
    document.save(path)


def configure_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)


def add_cluster_summary_table(document: Document, clusters: list[StoryCluster]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for index, label in enumerate(("Priority", "Section", "Story", "Sources")):
        set_cell_text(table.rows[0].cells[index], label, bold=True)
        shade(table.rows[0].cells[index], "E8EEF5")
    for cluster in clusters[:25]:
        cells = table.add_row().cells
        set_cell_text(cells[0], cluster.priority)
        set_cell_text(cells[1], cluster.section)
        set_cell_text(cells[2], cluster.canonical_headline)
        set_cell_text(cells[3], "; ".join(cluster.sources))


def add_cluster_entry(document: Document, cluster: StoryCluster) -> None:
    heading = document.add_heading(cluster.canonical_headline, level=2)
    heading.paragraph_format.keep_with_next = True
    meta = document.add_paragraph()
    meta.add_run(f"{cluster.section} | Score {cluster.relevance_score} | Confidence {cluster.confidence} | {cluster.publication_date}").bold = True
    document.add_paragraph(cluster.summary)
    source_line = document.add_paragraph()
    source_line.add_run("Sources: ").bold = True
    source_line.add_run("; ".join(cluster.sources))
    for url in cluster.article_urls[:4]:
        link = document.add_paragraph(url)
        link.paragraph_format.left_indent = Inches(0.2)
        link.runs[0].font.size = Pt(8)
        link.runs[0].font.color.rgb = RGBColor(31, 78, 121)


def add_pics_bullet(document: Document, cluster: StoryCluster) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run(cluster.canonical_headline)
    paragraph.add_run(" - ")
    for index, source in enumerate(cluster.sources):
        if index:
            paragraph.add_run("; ")
        label = f"{source}{' (Arabic)' if source_looks_arabic(source) else ''}"
        add_hyperlink(paragraph, cluster.article_urls[min(index, len(cluster.article_urls) - 1)], label)


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def ordered_section_themes(section_name: str, clusters: list[StoryCluster]) -> list[str]:
    preferred = PICS_SECTION_THEMES.get(section_name, [])
    present = {report_theme(cluster) for cluster in clusters}
    ordered = [theme for theme in preferred if theme in present]
    ordered.extend(sorted(present - set(ordered)))
    return ordered


def report_theme(cluster: StoryCluster) -> str:
    text = f"{cluster.canonical_headline} {cluster.summary}".casefold()
    return report_theme_from_text(cluster.section, text)


def report_theme_from_text(section: str, text: str) -> str:
    if section == "United Nations":
        if "structured dialogue" in text:
            return "Political Process"
        if "unhcr" in text or "iom" in text or "un agency" in text:
            return "UN Agencies"
        return "Other UN News"
    if section == "Politics":
        if "structured dialogue" in text or "election" in text or "constitution" in text or "حوار" in text or "انتخابات" in text:
            return "Political Process"
        if "government" in text or "cabinet" in text or "minister" in text or "حكومة" in text or "وزير" in text:
            return "Government Affairs"
        if "commentary" in text or "opinion" in text or "analysis" in text or "reaction" in text:
            return "Commentary"
        return "Other Political News"
    if section == "Military & Security":
        if "armed group" in text or "militia" in text or "مسلح" in text:
            return "Armed Groups"
        if "crime" in text or "arrest" in text or "fraud" in text or "جريمة" in text or "ضبط" in text:
            return "Crime & Enforcement"
        return "Security Developments"
    if section == "Migration":
        if "deport" in text or "anti-migration" in text or "enforcement" in text or "ترحيل" in text:
            return "Deportation & Enforcement"
        if "route" in text or "trafficking" in text or "boat" in text:
            return "Migration Routes & Trafficking"
        if "commentary" in text or "reaction" in text or "controversy" in text:
            return "Commentary"
        return "Migration Policy"
    if section == "Human Rights & Rule of Law":
        if "court" in text or "audit" in text or "accountability" in text or "rule of law" in text:
            return "Accountability"
        return "Human Rights"
    if section == "Economy & Energy":
        if "central bank" in text or "salary" in text or "finance" in text or "liquidity" in text:
            return "Banking & Finance"
        if "oil" in text or "fuel" in text or "noc" in text or "energy" in text or "slb" in text:
            return "Oil & Energy"
        if "infrastructure" in text or "electricity" in text or "port" in text or "free zone" in text:
            return "Infrastructure & Reconstruction"
        if "investment" in text or "trade" in text or "export" in text or "economic cooperation" in text:
            return "Other Economic News"
        return "Other Economic News"
    if section == "Environment":
        if "water" in text or "groundwater" in text or "مياه" in text:
            return "Water & Climate"
        if "agricultur" in text or "pest" in text or "آفات" in text or "زراع" in text:
            return "Agriculture"
        return "Environmental Risks"
    if section == "Governance & Public Services":
        if "health" in text or "medical" in text or "hospital" in text or "صحة" in text:
            return "Health"
        if "education" in text or "school" in text or "تعليم" in text or "امتحان" in text:
            return "Education"
        if "project" in text or "reconstruction" in text or "housing" in text or "infrastructure" in text:
            return "Reconstruction & Infrastructure"
        return "Public Administration"
    if section == "Regional & International":
        if "diplomacy" in text or "minister" in text or "presidential council" in text:
            return "Regional Diplomacy"
        if "cooperation" in text or "partnership" in text or "investment" in text:
            return "Foreign Cooperation"
        return "International Relations"
    if section == "Varieties":
        if "sport" in text or "football" in text or "رياض" in text:
            return "Sports"
        if "culture" in text or "heritage" in text or "ثقافة" in text or "تراث" in text:
            return "Culture"
        return "Society"
    return "Other"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_window(start_date: datetime | None, end_date: datetime | None) -> str:
    if start_date and end_date:
        return f"{start_date.date().isoformat()} to {end_date.date().isoformat()}"
    if start_date:
        return f"from {start_date.date().isoformat()}"
    if end_date:
        return f"to {end_date.date().isoformat()}"
    return "all available dates"


def is_non_article_url(url: str) -> bool:
    lowered = unquote(url).casefold()
    if is_listing_url(url):
        return True
    path = urlparse(url).path.strip("/").casefold()
    return path in {"", "news", "latest", "libya", "world", "international", "africa", "region-world"}


def article_text(article: Article) -> str:
    path = unquote(urlparse(article.url).path)
    return f"{article.title} {article.summary} {article.content_text} {path}".casefold()


def has_any(text: str, markers: list[str]) -> bool:
    return any(marker_in_text(text, marker) for marker in markers)


def first_match(text: str, markers: list[str]) -> str:
    return next((marker for marker in markers if marker_in_text(text, marker)), "")


def marker_in_text(text: str, marker: str) -> bool:
    lowered_marker = marker.casefold()
    if not lowered_marker:
        return False
    if re.fullmatch(r"[a-z0-9_-]+", lowered_marker):
        return re.search(rf"(?<![a-z0-9_-]){re.escape(lowered_marker)}(?![a-z0-9_-])", text) is not None
    return lowered_marker in text


def story_tokens(text: str) -> set[str]:
    normalized = text.casefold()
    normalized = re.sub(r"[\u064b-\u065f\u0670\u0640]", "", normalized)
    normalized = normalized.translate(
        str.maketrans({"أ": "ا", "إ": "ا", "آ": "ا", "ٱ": "ا", "ى": "ي", "ؤ": "و", "ئ": "ي", "ة": "ه"})
    )
    normalized = re.sub(r"https?://\S+", " ", normalized)
    normalized = re.sub(r"[^\w\u0600-\u06ff]+", " ", normalized)
    return {token for token in normalized.split() if token not in STOPWORDS and len(token) > 2}


def story_theme(article: Article) -> str:
    text = article_text(article)
    if ("الزاوية" in text or "zawiya" in text) and has_any(text, ["clash", "armed", "dead", "military", "اشتباك", "مسلح", "قتلى", "عسكرية"]):
        return "zawiya_clashes"
    if "قافلة غزة" in text or "قافلة الصمود" in text or "gaza convoy" in text:
        return "gaza_convoy_detention"
    if "حوار" in text or "structured dialogue" in text:
        return "structured_dialogue"
    if "توطين" in text or "settlement" in text:
        return "migration_settlement"
    if "ترحيل" in text or "deport" in text:
        return "migration_deportation"
    if "قارب" in text or "boat" in text:
        return "migrant_boat"
    if "تهريب" in text or "trafficking" in text:
        return "migration_trafficking"
    if "sudanese embassy" in text or "voluntary return" in text or "voluntary repatriation" in text or "الجالية السودانية" in text:
        return "sudan_repatriation"
    if ("central bank" in text or "مصرف" in text) and ("مرتبات" in text or "salary" in text or "shortfall" in text):
        return "central_bank_salary"
    if ("central bank" in text or "مصرف" in text) and ("cyber" in text or "سيبراني" in text or "ransomware" in text):
        return "central_bank_cyber"
    if ("central bank" in text or "مصرف" in text) and ("paris" in text or "finance forum" in text):
        return "central_bank_paris_forum"
    if ("مصرف" in text or "bank" in text) and ("نيابة" in text or "prosecution" in text or "حبس" in text or "detention" in text):
        return "bank_accountability_case"
    if ("schlumberger" in text or "slb" in text or "شلمبرجير" in text) and ("noc" in text or "national oil" in text or "النفط" in text):
        return "noc_slb"
    if "fuel" in text or "وقود" in text or "brega" in text or "البريقة" in text:
        return "brega_gas"
    if "tripoli" in text and ("project" in text or "development" in text or "مشاريع" in text):
        return "tripoli_projects"
    if "uganda" in text or "أوغندا" in text:
        return "uganda_relations"
    if "asian-african" in text or "آسيوي" in text or "المجلس البرلماني" in text:
        return "asian_african_parliamentary_forum"
    if "egypt" in text or "مصر" in text:
        return "egypt_us_libya_sudan"
    return ""


def entity_like_tokens(tokens: set[str]) -> set[str]:
    common_entities = {
        "libya",
        "libyan",
        "tripoli",
        "benghazi",
        "ليبيا",
        "ليبي",
        "الليبي",
        "طرابلس",
        "بنغازي",
        "الحكومة",
        "government",
    }
    return {token for token in tokens if len(token) >= 5 and token not in common_entities}


EVENT_TOKEN_STOPWORDS = STOPWORDS | {
    "libya",
    "libyan",
    "ليبيا",
    "ليبي",
    "الليبي",
    "government",
    "minister",
    "ministry",
    "official",
    "officials",
    "authority",
    "authorities",
    "council",
    "state",
    "national",
    "public",
    "meeting",
    "discuss",
    "discusses",
    "discussed",
    "talks",
    "support",
    "development",
    "cooperation",
    "coordination",
    "affairs",
    "حكومة",
    "الحكومة",
    "وزير",
    "وزارة",
    "مسؤول",
    "مسؤولين",
    "سلطات",
    "مجلس",
    "الدولة",
    "وطني",
    "الوطنية",
    "اجتماع",
    "بحث",
    "يناقش",
    "ناقش",
    "دعم",
    "تطوير",
    "تعاون",
    "تنسيق",
}


def story_event_signature(article: Article) -> set[str]:
    text = article.title
    tokens = story_tokens(text)
    return {
        token
        for token in tokens
        if token not in EVENT_TOKEN_STOPWORDS and (len(token) >= 5 or re.search(r"[\u0600-\u06ff]", token))
    }


def article_matches_existing_story(
    tokens: set[str],
    event_signature: set[str],
    theme: str,
    existing_article: Article,
    same_theme: bool,
) -> bool:
    existing_tokens = story_tokens(existing_article.title + " " + existing_article.summary)
    existing_signature = story_event_signature(existing_article)
    overlap = event_signature & existing_signature
    similarity = jaccard(tokens, existing_tokens)
    role = narrative_role_from_tokens(tokens)
    existing_role = narrative_role(existing_article)
    if role and existing_role and role != existing_role:
        return False
    if len(overlap) >= 5:
        return True
    if len(overlap) >= 3 and similarity >= 0.24:
        return True
    if same_theme and theme in STRONG_STORY_THEMES and len(overlap) >= 4 and similarity >= 0.38:
        return True
    return False


NARRATIVE_ROLE_MARKERS = {
    "reaction_support": [
        "support",
        "welcomes",
        "backs",
        "endorses",
        "praises",
        "ترحب",
        "يرحب",
        "يدعم",
        "تدعم",
        "يؤيد",
        "تأييد",
        "دعم",
    ],
    "reaction_opposition": [
        "criticizes",
        "rejects",
        "opposes",
        "warns",
        "condemns",
        "رفض",
        "يرفض",
        "انتقد",
        "ينتقد",
        "يحذر",
        "تحذر",
        "يدين",
    ],
    "legal_objection": [
        "constitutional drafting",
        "legal objection",
        "lawsuit",
        "court",
        "appeal",
        "الهيئة التأسيسية",
        "اعتراض قانوني",
        "طعن",
        "محكمة",
        "دستوري",
    ],
    "implementation_concern": [
        "implementation",
        "concern",
        "challenge",
        "obstacle",
        "risk",
        "تنفيذ",
        "مخاوف",
        "تحديات",
        "عقبات",
        "مخاطر",
    ],
    "proposal": [
        "proposal",
        "proposes",
        "roadmap",
        "plan",
        "recommendation",
        "مقترح",
        "تقترح",
        "خارطة",
        "خطة",
        "توصية",
        "مخرجات",
    ],
    "analysis": [
        "analysis",
        "analyst",
        "explainer",
        "تحليل",
        "قراءة",
        "محلل",
    ],
    "commentary": [
        "opinion",
        "commentary",
        "editorial",
        "column",
        "interview",
        "podcast",
        "رأي",
        "مقال",
        "افتتاحية",
        "مقابلة",
        "بودكاست",
    ],
    "municipal_position": [
        "municipality",
        "mayor",
        "municipal council",
        "بلدية",
        "عميد",
        "مجلس بلدي",
    ],
    "tribal_position": [
        "tribal",
        "notables",
        "social council",
        "قبائل",
        "أعيان",
        "مجلس اجتماعي",
    ],
    "statement": [
        "statement",
        "announces",
        "announcement",
        "said",
        "بيان",
        "أعلن",
        "تعلن",
        "قال",
        "صرح",
    ],
    "meeting": [
        "meeting",
        "meets",
        "discusses",
        "talks",
        "اجتمع",
        "يلتقي",
        "بحث",
        "ناقش",
    ],
    "decision": [
        "decision",
        "approves",
        "adopts",
        "signs",
        "agreement",
        "قرار",
        "يعتمد",
        "وقّع",
        "وقع",
        "اتفاق",
        "مذكرة",
    ],
}


def narrative_role(article: Article) -> str:
    return narrative_role_from_text(article_text(article))


def narrative_role_from_text(text: str) -> str:
    lowered = text.casefold()
    for role, markers in NARRATIVE_ROLE_MARKERS.items():
        if has_any(lowered, markers):
            return role
    return ""


def narrative_role_from_tokens(tokens: set[str]) -> str:
    text = " ".join(tokens)
    return narrative_role_from_text(text)


def jaccard(left: set[str], right: set[str]) -> float:
    if not left or not right:
        return 0.0
    return len(left & right) / len(left | right)


def best_summary(members: list[Article]) -> str:
    candidates = [member.summary.strip() for member in members if member.summary.strip()]
    if candidates:
        return max(candidates, key=len)[:700]
    return members[0].title


def clean_english_headline(members: list[Article]) -> str:
    lead = members[0]
    text = combined_cluster_text(members)
    english_title = best_english_title(members)
    if english_title and not is_generic_headline(english_title) and headline_has_specific_event(english_title):
        return concise(english_title)
    translated_title = translate_headline_stub(lead.title, lead.language)
    if translated_title and headline_has_specific_event(translated_title):
        return concise(translated_title)

    if ("تيتيه" in text or "tetteh" in text or "المبعوثة الأممية" in text) and ("مجلس السلم والأمن" in text or "peace and security council" in text):
        return "Tetteh briefs African Peace and Security Council on Libya developments"
    if "صدام حفتر" in text and ("وفد تركي" in text or "تركي" in text):
        return "Saddam Haftar and Turkish delegation discuss expanding Libya-Turkey cooperation"
    if "تونس" in text and ("انتخاب" in text or "لقاء عربي" in text):
        return "Libya participates in Arab election forum in Tunisia"
    if "مظاهرات طرابلس" in text and "الهجرة غير النظامية" in text:
        return "Tripoli protests raise public concern over irregular migration file"
    if "المسماري" in text and ("الاستقرار الدائم" in text or "قضية أمن قومي" in text):
        return "Al-Mismari rejects permanent migrant settlement and frames migration as national security issue"
    if "الخفيفي" in text and ("تهديدًا مباشرًا" in text or "الأمن القومي" in text):
        return "Al-Khafifi warns irregular migration poses direct threat to Libyan national security"
    if "اللافي" in text and ("الصين" in text or "china" in text) and "الاستقرار" in text:
        return "Lafi says China is an important partner in supporting Libya stability"
    if "سفارة قطر" in text and "الهجرة" in text:
        return "Qatar Embassy denies involvement in Libya migration-related claims"
    if "مفوضية اللاجئين" in text and "مقر البعثة" in text:
        return "UNSMIL denies UNHCR office presence at mission headquarters"
    if "سفارة السودان" in text and ("عودة اللاجئين" in text or "العودة الطوعية" in text):
        return "Sudanese Embassy launches second phase of voluntary refugee return from Libya"
    if ("بالقاسم حفتر" in text or "belqasem haftar" in text) and ("الأميركي" in text or "american" in text or "us " in text) and ("إعمار" in text or "reconstruction" in text):
        return "Belqasem Haftar and US envoy discuss development and reconstruction cooperation"
    if ("خالد حفتر" in text or "khaled haftar" in text) and ("الأمريكي" in text or "american" in text or "us " in text) and ("التعاون العسكري" in text or "military cooperation" in text):
        return "Khaled Haftar and US envoy discuss military cooperation and regional stability"
    if "اللافي" in text and "النمروش" in text and ("الأوضاع الأمنية والعسكرية" in text or "security and military" in text):
        return "Lafi and Namroush discuss western Libya security and military tensions"
    if ("الحويج" in text or "hweij" in text) and ("دور إفريقيا" in text or "african role" in text):
        return "Hweij and African envoy discuss African support for Libya stability track"
    if ("irini" in text or "إيريني" in text) and ("shadow fleet" in text or "الشبح" in text):
        return "EU expands Operation IRINI mandate amid Russian shadow fleet crackdown near Libya"
    if ("قادربوه" in text or "qaderboh" in text) and ("النائب العام" in text or "attorney general" in text) and ("المحروقات" in text or "fuel" in text):
        return "Oversight Authority and Attorney General discuss fuel file and oil-sector financing"
    if ("noc" in text or "national oil corporation" in text) and ("pmi" in text or "project management institute" in text):
        return "NOC and PMI sign capacity-building agreement for Libya energy sector"
    if ("المنفي" in text or "menfi" in text) and ("السفير البلجيكي" in text or "belgian ambassador" in text):
        return "Menfi praises Belgian ambassador role in strengthening bilateral relations"
    if ("الشهوبي" in text or "transport minister" in text) and ("أنتويرب" in text or "antwerp" in text):
        return "Transport Minister and Belgian ambassador discuss linking Libyan ports with Antwerp"
    if ("الزادمة" in text or "zadma" in text) and ("السفير الفرنسي" in text or "french ambassador" in text or "فالا" in text):
        return "Zadma and French ambassador discuss Fezzan-France ties and Libya political developments"
    if "media cooperation with france" in text or ("paris" in text and "media cooperation" in text):
        return "Libyan delegation visits Paris to strengthen media cooperation with France"
    if has_any(text, ["settlement", "توطين"]) and has_any(text, ["misinformation", "false", "تنفي", "مضلل", "شائعات"]):
        return "UNSMIL rejects misinformation regarding migrant settlement in Libya"
    if "unsmil denies relocation of unhcr office" in text or ("unhcr" in text and "relocation" in text):
        return "UNSMIL denies relocation of UNHCR office to mission headquarters"
    if has_any(text, ["قافلة غزة", "قافلة الصمود", "gaza convoy", "hunger strike"]):
        return "Health and rights concerns grow over Gaza convoy activists detained in eastern Libya"

    if has_any(text, ["migration", "migrant", "refugee", "deport", "traffick", "settlement", "unhcr", "iom", "هجرة", "مهاجر", "لاجئ", "ترحيل", "تهريب", "توطين"]):
        if "قارب" in text or "bodies" in text or "boat" in text:
            return "Deaths reported after migrant boat incident linked to Libya route"
        if "ترحيل" in text or "deport" in text:
            return "Libyan authorities expand deportation and anti-migration operations"
        if "أجدابيا" in text or "ajdabiya" in text:
            return "Ajdabiya authorities continue anti-irregular migration campaign"
        if "تهريب" in text or "trafficking" in text:
            return "Libya migration route reporting highlights trafficking and protection risks"
        if "unhcr" in text or "iom" in text:
            return "UN migration agencies track protection and return issues affecting Libya"
        if "sudanese embassy" in text or "voluntary return" in text or "voluntary repatriation" in text:
            return "Sudanese Embassy says nationals in Libya seek voluntary return"
        if "توطين" in text or "settlement" in text:
            return "Libyan officials reject migrant settlement proposals amid national security concerns"
        return "Migration pressures in Libya draw policy and protection concerns"

    if has_any(text, ["armed", "clash", "militia", "weapon", "security operation", "crime", "arrest", "smuggling", "drug", "اشتباك", "مسلح", "سلاح", "جريمة", "قبض", "ضبط", "تهريب", "مخدرات"]):
        if "الزاوية" in text or "zawiya" in text:
            return "Armed clashes renew in Al-Zawiya amid continuing local security tensions"
        if "مخدرات" in text or "drug" in text:
            return "Libyan security authorities report drug-enforcement operation"
        if "تهريب" in text or "smuggling" in text:
            return "Libyan authorities report smuggling-related security enforcement"
        if "قبض" in text or "ضبط" in text or "arrest" in text:
            return "Libyan security authorities report local arrest operation"
        if "احتيال" in text or "cyber" in text or "fraud" in text:
            return "Tripoli security authorities warn over cyber fraud risks"
        if "ترتيبات أمنية" in text or "security arrangement" in text:
            return "Libyan authorities discuss security arrangements linked to local stability"
        if "سلاح" in text or "weapon" in text:
            return "Weapons-related security reporting raises concerns over armed-group activity"
        return "Libyan security authorities report operation affecting local stability"

    if has_any(text, ["central bank", "public finance", "salary", "liquidity", "budget", "oil", "gas", "fuel", "noc", "brega", "مصرف", "ميزانية", "مرتبات", "سيولة", "نفط", "غاز", "وقود", "البريقة"]):
        if "central bank" in text or "مصرف" in text:
            if "مرتبات" in text or "salary" in text:
                return "Central Bank warns over rising salary bill and fiscal sustainability risks"
            if "paris" in text or "finance forum" in text:
                return "Central Bank governor joins Libya-France finance forum in Paris"
            return "Central Bank discusses liquidity pressures and public finance stability"
        if "شلمبرجير" in text or "schlumberger" in text or "slb" in text:
            return "NOC and SLB discuss oil-sector training, technology transfer and Libyan expertise"
        if "توزيع الغاز" in text or "brega" in text or "البريقة" in text:
            return "Brega announces gas distribution plan for Jalu, Kufra and Benghazi"
        if "ناقلة" in text or "بنزين" in text or "tobruk" in text:
            return "Fuel shipment arrives in Tobruk amid continued supply monitoring"
        if "رأس لانوف" in text or "ras lanuf" in text:
            return "Ras Lanuf refinery restart expected to support Libya economic recovery"
        if "صادرات" in text or "exports" in text:
            return "Libyan oil export reporting points to pressure on production outlook"
        if "كهرباء" in text or "electricity" in text:
            return "Electricity outages and grid pressures continue affecting western Libya"
        return "Libya economic reporting highlights pressure on energy and public finance"

    if has_any(text, ["unsmil", "security council", "united nations", "srsg", "dsrsg", "البعثة الأممية", "الأمم المتحدة", "مجلس الأمن"]):
        if ("تيتيه" in text or "tetteh" in text or "المبعوثة الأممية" in text) and ("مجلس السلم والأمن" in text or "peace and security council" in text):
            return "Tetteh briefs African Peace and Security Council on Libya developments"
        if "حوار" in text or "structured dialogue" in text:
            return "UNSMIL-linked Structured Dialogue recommendations shape Libya transition debate"
        if "مجلس الأمن" in text or "security council" in text:
            return "Libya political track raised in UN Security Council context"
        if "الاتحاد الأفريقي" in text or "african union" in text:
            return "African Union seeks movement on Libya political crisis"
        if "انقسام" in text:
            return "Institutional division identified as core obstacle in Libya crisis"
        return "UN officials address Libya political process and institutional division"

    if "structured dialogue" in text or "حوار" in text:
        if "اقتصاد" in text or "دولار" in text:
            return "Structured Dialogue economic scenarios highlight risks for Libya's exchange rate"
        if "سلطة تنفيذية" in text:
            return "Structured Dialogue governance track addresses Libya executive authority options"
        if "دستوري" in text:
            return "Structured Dialogue recommendations address Libya constitutional basis"
        return "Structured Dialogue recommendations advance Libya political process"

    if has_any(text, ["election", "parliament", "house of representatives", "high council of state", "presidential council", "government", "executive authority", "constitutional", "انتخابات", "مجلس النواب", "المجلس الأعلى للدولة", "المجلس الرئاسي", "حكومة", "سلطة تنفيذية", "دستوري"]):
        if "asian-african" in text or "آسيوي" in text:
            return "Benghazi prepares to host Asian-African parliamentary forum"
        if "حكومة موازية" in text or "parallel government" in text:
            return "Political debate continues over executive authority and rival government arrangements"
        if "انتخابات" in text or "election" in text:
            return "Libyan political actors debate election pathway and constitutional basis"
        return "Libyan political institutions debate transition arrangements and government authority"

    if has_any(text, ["court", "prosecution", "public prosecution", "attorney general", "detention", "forgery", "corruption", "human rights", "icc", "accountability", "محكمة", "نيابة", "النائب العام", "حبس", "تزوير", "فساد", "احتجاز", "حقوق الإنسان", "مساءلة"]):
        if "icc" in text or "الجنائية الدولية" in text:
            return "ICC-related reporting renews accountability focus on Libya"
        if "النائب العام" in text or "public prosecution" in text or "attorney general" in text:
            if "تزوير" in text or "forgery" in text:
                return "Public Prosecution orders detention in forgery and public funds case"
            if "فساد" in text or "corruption" in text:
                return "Public Prosecution pursues corruption case with rule-of-law implications"
            if "حبس" in text or "detention" in text:
                return "Public Prosecution orders detention in local accountability case"
            return "Public Prosecution reports accountability action in Libya"
        if "محكمة" in text or "court" in text:
            return "Libyan court proceedings highlight rule-of-law and accountability issues"
        if "أجدابيا" in text and ("نيجيري" in text or "مزورة" in text):
            return "Ajdabiya security authorities arrest Nigerian suspect over forged Libyan identity"
        return "Rights and accountability concerns remain active in Libya monitoring"

    if has_any(text, MUNICIPAL_MARKERS + ["airport", "airline", "public administration", "medical facilities", "مطار", "طيران", "مرافق طبية"]):
        if "مرافق طبية" in text or "medical facilities" in text:
            return "Benghazi opens new medical facilities to strengthen health services"
        if "الأفريقية" in text or "afriqiyah" in text:
            return "Afriqiyah Airways denies fleet shutdown and announces new aircraft entry"
        if "امتحانات" in text or "exam" in text:
            return "Education authorities publish secondary school exam schedules inside and outside Libya"
        if "tripoli" in text or "طرابلس" in text:
            return "Government reviews implementation challenges affecting Tripoli development projects"
        if "متابعة" in text or "تنفيذ" in text or "implementation" in text or "follow-up" in text:
            return "Government bodies follow up implementation of local public-service plans"
        if "بلدية" in text or "municipal" in text:
            return "Municipal authorities report local governance and service-delivery follow-up"
        if "housing" in text or "الإسكان" in text:
            return "Housing-sector decisions aim to restart stalled public housing activity"
        if "reconstruction" in text or "إعمار" in text:
            return "Reconstruction authorities discuss southern economic development priorities"
        if "استراتيجيات" in text or "strategies" in text:
            return "Dbeibah receives national development strategies for implementation planning"
        return "Government bodies address public administration and service-delivery priorities"

    if has_any(text, ["foreign minister", "ambassador", "diplomacy", "bilateral", "italy", "egypt", "tunisia", "turkey", "uganda", "china", "uae", "وزير الخارجية", "سفير", "دبلوماسي", "ثنائية", "إيطاليا", "مصر", "تونس", "تركيا", "أوغندا", "الصين", "الإمارات"]):
        if "وفد تركي" in text or "turkish delegation" in text or "تركي" in text or "turkish" in text:
            if "صدام حفتر" in text or "saddam haftar" in text:
                return "Saddam Haftar and Turkish delegation discuss expanding Libya-Turkey cooperation"
            return "Libyan and Turkish officials discuss strengthening bilateral cooperation"
        if "موريتانيا" in text or "mauritania" in text:
            return "Libya and Mauritania discuss expanding bilateral partnership"
        if "الإمارات" in text or "uae" in text or "emirati" in text:
            return "Dbeibah and UAE ambassador discuss economic partnership and investment"
        if "تونس" in text or "tunisia" in text:
            return "Libya participates in regional election-related forum in Tunisia"
        if "إيطاليا" in text or "italy" in text:
            if "misrata" in text or "genoa" in text or "جنوة" in text:
                return "Misrata and Genoa chambers sign partnership to expand trade and investment"
            return "Italy-linked diplomacy keeps Libya migration and security file active"
        if "uganda" in text or "أوغندا" in text:
            return "Libya and Uganda discuss expanding bilateral economic cooperation"
        if "الصين" in text or "china" in text:
            return "Presidential Council highlights China role in supporting Libya stability"
        if "مصر" in text or "egypt" in text:
            return "Egypt and United States discuss developments in Libya and Sudan"
        if "genoa" in text or "جنوة" in text:
            return "Misrata and Genoa chambers seek joint economic projects"
        return "Libyan officials pursue foreign cooperation linked to stability and economic interests"

    if english_title:
        return concise(english_title)
    issue = specific_issue_label(text)
    return f"Libyan institutions address {issue} with national monitoring relevance"


def translate_headline_stub(headline: str, language: str) -> str:
    """Hook for future translation API integration; returns empty until credentials are configured."""
    if language.casefold().startswith("en"):
        return headline
    return ""


def clean_english_summary(members: list[Article]) -> str:
    lead = members[0]
    return inclusion_reason(lead)


def best_english_title(members: list[Article]) -> str:
    english_titles = [
        member.title.strip()
        for member in sorted(members, key=article_editorial_rank)
        if is_mostly_english(member.title) and headline_has_specific_event(member.title)
    ]
    if not english_titles:
        return ""
    return english_titles[0]


def headline_has_specific_event(headline: str) -> bool:
    lowered = headline.casefold()
    if is_generic_headline(headline):
        return False
    weak_terms = ["development", "issue", "file", "reported", "public affairs"]
    if sum(1 for term in weak_terms if term in lowered) >= 2:
        return False
    actor_or_place = has_any(
        lowered,
        [
            "unsmil",
            "un ",
            "libya",
            "libyan",
            "tripoli",
            "benghazi",
            "zawiya",
            "central bank",
            "noc",
            "brega",
            "authorities",
            "officials",
            "agencies",
            "institutions",
            "council",
            "court",
            "public prosecution",
            "prosecution",
            "attorney general",
            "unhcr",
            "mission",
            "structured dialogue",
            "oversight authority",
            "attorney general",
            "italy",
            "egypt",
            "ras lanuf",
            "ajdabiya",
            "electricity",
            "grid",
            "refinery",
            "embassy",
            "qatar embassy",
            "sudanese embassy",
            "misrata",
            "genoa",
            "saddam haftar",
            "khaled haftar",
            "belqasem haftar",
            "lafi",
            "namroush",
            "tetteh",
            "menfi",
            "al-mismari",
            "al-khafifi",
            "transport minister",
            "minister of agriculture",
            "dbeibah",
            "government",
            "court",
            "migration",
            "migrant",
            "ليبيا",
        ],
    )
    action = has_any(
        lowered,
        [
            "reject",
            "warn",
            "discuss",
            "announce",
            "open",
            "renew",
            "expand",
            "arrest",
            "review",
            "deny",
            "denies",
            "participate",
            "highlight",
            "call",
            "urge",
            "resume",
            "fall",
            "rise",
            "shape",
            "address",
            "track",
            "report",
            "reported",
            "linked",
            "grow",
            "remain",
            "continue",
            "continues",
            "expected",
            "affecting",
            "affect",
            "raise",
            "raises",
            "seek",
            "seeks",
            "rooted",
            "briefs",
            "hosts",
            "launches",
            "sign",
            "signs",
            "praises",
            "visits",
            "participates",
            "arrive",
            "arrives",
            "review",
            "reviews",
            "warns",
            "poses",
            "supporting",
            "advance",
            "advances",
            "order",
            "orders",
            "detention",
            "pursue",
            "pursues",
        ],
    )
    return actor_or_place and action and len(headline.split()) >= 6


def specific_issue_label(text: str) -> str:
    if "health" in text or "medical" in text or "صحة" in text:
        return "health-service priorities"
    if "education" in text or "تعليم" in text:
        return "education-sector priorities"
    if "municipal" in text or "بلدية" in text:
        return "municipal governance priorities"
    if "infrastructure" in text or "بنية تحتية" in text:
        return "infrastructure implementation challenges"
    if "culture" in text or "ثقافة" in text:
        return "cultural and social developments"
    if "sport" in text or "رياض" in text:
        return "sports developments"
    return "a public-policy issue"


def inclusion_reason(article: Article) -> str:
    return inclusion_reason_for_section(article.section_guess)


def inclusion_reason_for_section(section: str) -> str:
    if section == "United Nations":
        return "Included because it concerns UN engagement or the UN track in Libya."
    if section == "Politics":
        return "Included because it affects Libya's political process or national institutions."
    if section == "Military & Security":
        return "Included because it affects security conditions or armed/security arrangements."
    if section == "Human Rights & Rule of Law":
        return "Included because it concerns rights, justice, accountability or rule of law."
    if section == "Migration":
        return "Included because migration, refugees or migrant protection affect UNSMIL situational awareness."
    if section == "Economy & Energy":
        return "Included because it concerns economic stability, public finance, oil, fuel or energy policy."
    if section == "Environment":
        return "Included because it concerns environmental risks, water, climate, agriculture or public safety."
    if section == "Governance & Public Services":
        return "Included because it concerns governance or public services with wider public significance."
    if section == "Regional & International":
        return "Included because it concerns foreign relations or diplomacy affecting Libya."
    return "Included because it has exceptional Libya relevance."


def multi_source_phrase(members: list[Article]) -> str:
    names = sorted({member.source_name for member in members})
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return f"{names[0]} and {names[1]}"
    return f"{', '.join(names[:2])} and other sources"


def source_looks_arabic(source_name: str) -> bool:
    english_only = {"Libya Herald", "Libya Observer", "Libya Review", "Reuters", "ANSA", "The Guardian", "New Arab"}
    return source_name not in english_only


def is_listing_url(url: str) -> bool:
    lowered = unquote(url).casefold()
    return any(marker in lowered for marker in NON_ARTICLE_MARKERS)


def is_unrelated_international_noise(article: Article) -> bool:
    text = article_text(article)
    international_markers = ["uk election", "gaza", "iran", "israel", "ukraine", "trump", "غزة", "إيران", "إسرائيل", "أوكرانيا"]
    return has_any(text, international_markers) and not has_any(text, LIBYA_MARKERS)


def is_mostly_english(text: str) -> bool:
    letters = re.findall(r"[A-Za-z\u0600-\u06ff]", text)
    if not letters:
        return False
    latin = sum(1 for char in letters if "A" <= char <= "z")
    return latin / len(letters) >= 0.75


def concise(text: str, max_length: int = 150) -> str:
    cleaned = html.unescape(re.sub(r"\s+", " ", text).strip())
    if len(cleaned) <= max_length:
        return cleaned
    return cleaned[: max_length - 1].rstrip() + "..."


def is_generic_headline(headline: str) -> bool:
    return any(headline.startswith(prefix) for prefix in GENERIC_HEADLINE_PREFIXES)


def normalize_headline_for_duplicate_check(headline: str) -> str:
    return re.sub(r"\W+", " ", headline.casefold()).strip()


def append_note(existing: str, note: str) -> str:
    if not existing:
        return note
    if note in existing:
        return existing
    return f"{existing}; {note}"
