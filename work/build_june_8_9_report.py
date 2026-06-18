from __future__ import annotations

import csv
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
REPORTS = OUTPUT / "reports"
REPORTS.mkdir(parents=True, exist_ok=True)

APPROVED = OUTPUT / "approved_headlines.csv"
REVIEW = OUTPUT / "review_queue.csv"
DEBUG = OUTPUT / "source_debug_report.csv"
REPORT = REPORTS / "UNSMIL_PICS_Libya_Media_Monitoring_2026-06-08_2026-06-09.docx"


SECTION_ORDER = [
    "United Nations",
    "Politics",
    "Governance",
    "Military & Security",
    "Migration",
    "Economy",
    "Municipalities & Public Services",
    "Health",
    "Human Rights",
    "Environment",
    "Regional & International",
    "Varieties",
]


def clean_row(row: dict[str, str]) -> dict[str, str]:
    return {key.lstrip("\ufeff"): value for key, value in row.items()}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return [clean_row(row) for row in csv.DictReader(handle)]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_metadata_table(doc: Document, approved: list[dict[str, str]], review: list[dict[str, str]]) -> None:
    dates = Counter(item["publication_date"] for item in approved)
    duplicates_removed = 2
    rows = [
        ("Monitoring window", "8-9 June 2026"),
        ("Approved headlines", str(len(approved))),
        ("Review queue", str(len(review))),
        ("Duplicates removed", str(duplicates_removed)),
        ("Date coverage", ", ".join(f"{day}: {count}" for day, count in sorted(dates.items()))),
        ("Generated", date.today().isoformat()),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for index, label in enumerate(("Field", "Value")):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, label, bold=True)
        set_cell_shading(cell, "E8EEF5")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)


def add_source_table(doc: Document, approved: list[dict[str, str]], review: list[dict[str, str]]) -> None:
    approved_by_source = Counter(item["source_name"] for item in approved)
    review_by_source = Counter(item["source_name"] for item in review)
    sources = sorted(set(approved_by_source) | set(review_by_source), key=lambda s: (-approved_by_source[s], s))
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, label in enumerate(("Source", "Approved", "Review")):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, label, bold=True)
        set_cell_shading(cell, "E8EEF5")
    for source in sources:
        if approved_by_source[source] == 0 and review_by_source[source] == 0:
            continue
        cells = table.add_row().cells
        set_cell_text(cells[0], source)
        set_cell_text(cells[1], str(approved_by_source[source]))
        set_cell_text(cells[2], str(review_by_source[source]))


def add_zero_source_note(doc: Document, debug_rows: list[dict[str, str]]) -> None:
    zero_rows = [row for row in debug_rows if row.get("accepted_count") == "0"]
    if not zero_rows:
        return
    doc.add_heading("Sources With No Approved Items", level=2)
    paragraph = doc.add_paragraph()
    paragraph.add_run("These sources were checked but produced no approved headlines for this draft: ").bold = True
    paragraph.add_run(
        "; ".join(
            f"{row['source_name']} ({row.get('zero_result_reason') or 'no reason recorded'})"
            for row in zero_rows
        )
        + "."
    )


def add_headline_sections(doc: Document, approved: list[dict[str, str]]) -> None:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for item in approved:
        grouped[item.get("section_guess") or "Other"].append(item)

    seen_sections = set()
    for section in SECTION_ORDER + sorted(grouped):
        if section in seen_sections or section not in grouped:
            continue
        seen_sections.add(section)
        items = sorted(grouped[section], key=lambda row: (row["publication_date"], row["source_name"], row["headline_original"]))
        doc.add_heading(f"{section} ({len(items)})", level=1)
        for item in items:
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(item["headline_original"])
            run.bold = True
            paragraph.add_run(f" — {item['source_name']}, {item['publication_date']}")
            if item.get("article_snippet"):
                snippet = doc.add_paragraph(item["article_snippet"])
                snippet.paragraph_format.left_indent = Inches(0.25)
                snippet.paragraph_format.space_after = Pt(2)
            link = doc.add_paragraph(item["article_url"])
            link.paragraph_format.left_indent = Inches(0.25)
            link.runs[0].font.size = Pt(8)
            link.runs[0].font.color.rgb = RGBColor(31, 78, 121)


def build() -> Path:
    approved = read_csv(APPROVED)
    review = read_csv(REVIEW)
    debug_rows = read_csv(DEBUG)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("UNSMIL/PICS Libya Media Monitoring")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Headline draft: 8-9 June 2026")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    note = doc.add_paragraph()
    note.add_run("Editorial status: ").bold = True
    note.add_run(
        "Draft built from approved scraper output only. Items with missing, ambiguous, or conflicting dates remain in the separate review queue."
    )

    doc.add_heading("Run Summary", level=1)
    add_metadata_table(doc, approved, review)

    doc.add_heading("Source Contribution", level=1)
    add_source_table(doc, approved, review)
    add_zero_source_note(doc, debug_rows)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Approved Headlines", level=1)
    add_headline_sections(doc, approved)

    doc.save(REPORT)
    return REPORT


if __name__ == "__main__":
    print(build())
